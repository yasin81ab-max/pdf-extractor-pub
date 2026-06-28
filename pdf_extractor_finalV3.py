# -*- coding: utf-8 -*-
"""
PDF Catalog Extractor
- ورودی: یک/چند فولدر یا یک/چند PDF تکی
- برای هر فولدر: ابتدا MERG.pdf برای کشف ویژگی‌ها (یک API call) → تأیید کاربر → استخراج بقیه PDFها
- هر PDF (غیر MERG) یک API call برای استخراج؛ rotation فقط هنگام خطا
- خروجی: برای هر فولدر یک Excel + یک Excel کلی نهایی
- لاگ ساده برای کاربر (مرحله + خلاصه خطا)
"""

import json
import re
import time
import logging
import sys
import os
import base64
import hashlib
import sqlite3
import uuid
from collections import deque
from datetime import datetime
from pathlib import Path
from threading import Lock, Thread, Event
from typing import Optional, Dict, List, Any, Callable

# ======================== EXE / PATH FIX ========================

def get_app_dir() -> Path:
    """
    مسیر واقعی کنار فایل exe (یا .py در حالت توسعه).
    با --onefile، sys.executable مسیر exe است، نه فولدر موقت.
    """
    if getattr(sys, 'frozen', False):
        return Path(sys.executable).parent
    return Path(__file__).parent


def get_resource_path(filename: str) -> Path:
    """
    مسیر فایل‌های داخلی بسته‌شده در exe (مثل آیکون).
    PyInstaller فایل‌های --add-data را در sys._MEIPASS قرار می‌دهد.
    """
    if getattr(sys, 'frozen', False):
        base = Path(sys._MEIPASS)
    else:
        base = Path(__file__).parent
    return base / filename


# ======================== ENCODING FIX (Windows + noconsole) ========================

if sys.platform == "win32":
    if sys.stdout is not None:
        try:
            import io as _io
            sys.stdout = _io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")
        except Exception:
            pass
    if sys.stderr is not None:
        try:
            import io as _io
            sys.stderr = _io.TextIOWrapper(sys.stderr.buffer, encoding="utf-8")
        except Exception:
            pass

# ======================== LOGGING SETUP ========================

APP_DIR = get_app_dir()
LOG_PATH = APP_DIR / "pdf_extractor.log"

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    handlers=[
        logging.FileHandler(str(LOG_PATH), encoding="utf-8"),
        *([] if (sys.stdout is None) else [logging.StreamHandler()]),
    ],
)

# ======================== IMPORTS ========================

import requests
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

try:
    from pypdf import PdfReader, PdfWriter
    PYPDF_AVAILABLE = True
except ImportError:
    try:
        from PyPDF2 import PdfReader, PdfWriter
        PYPDF_AVAILABLE = True
    except ImportError:
        PYPDF_AVAILABLE = False
        logging.warning("pypdf/PyPDF2 نصب نیست — نمونه‌برداری صفحات MERG غیرفعال است (کل فایل ارسال می‌شود).")

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx نصب نیست — خروجی Word غیرفعال است. نصب: pip install python-docx")

try:
    from PIL import Image
    try:
        from PIL import PngImagePlugin
    except Exception:
        PngImagePlugin = None
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False
    logging.warning("Pillow نصب نیست — درج توضیحات در متادیتای عکس غیرفعال است. نصب: pip install Pillow")

# ======================== CONSTANTS ========================

DEFAULT_CONFIG_NAME = "config.json"
MERG_FILENAME = "MERG.pdf"
MAX_RETRIES_PER_KEY = 3
RETRY_DELAY = 1
API_RATE_DELAY = 0.3
TIMEOUT_S = 120
TEMPERATURE = 0.1
MERG_SAMPLE_PAGE_CAP = 20

# لیست موضوع‌ها از این آدرس raw گیت‌هاب خوانده می‌شود (بدون نیاز به اعتبارنامه).
# کاربر/توسعه‌دهنده فقط همین فایل را در ریپو آپدیت می‌کند (یا با GitHub Action از شیت).
SUBJECTS_URL = "https://raw.githubusercontent.com/USERNAME/REPO/main/subjects.json"
SUBJECTS_FETCH_TIMEOUT = 8

# تنظیمات توضیح‌گذاری تصاویر (vision)
IMAGE_DESC_ENABLED_DEFAULT = True
IMAGE_DESC_BATCH = 12            # تعداد عکس در هر فراخوانی vision
IMAGE_TYPES = ["نمودار", "عکس کالا", "جدول", "اطلاعات تماس", "لوگو", "سایر"]

# ======================== GOOGLE DRIVE + TELEGRAM ========================
# ┌─────────────────────────────────────────────────────────────────────┐
# │  تنظیمات توسعه‌دهنده — فقط قبل از build تغییر دهید               │
# │  کاربر نهایی این بخش را نمی‌بیند و نیازی به دستکاری ندارد        │
# └─────────────────────────────────────────────────────────────────────┘

# نام فولدر کلی در Google Drive (زیر آن موضوع و تاریخ اجرا ساخته می‌شود)
DRIVE_ROOT_FOLDER_NAME = "CatalogExtractor"

# OAuth client credentials — از Google Cloud Console → APIs & Services → Credentials
# نوع: OAuth 2.0 Client ID → Desktop App
OAUTH_CLIENT_ID     = ""   # ← client_id را اینجا بگذارید
OAUTH_CLIENT_SECRET = ""   # ← client_secret را اینجا بگذارید
OAUTH_REDIRECT_URI  = "urn:ietf:wg:oauth:2.0:oob"   # تغییر ندهید
DRIVE_SCOPES        = ["https://www.googleapis.com/auth/drive"]

# Telegram — ربات اطلاع‌رسانی (فقط برای توسعه‌دهنده، کاربر خبر ندارد)
TGRAM_BOT_TOKEN = ""        # ← توکن ربات (از @BotFather)
TGRAM_CHAT_ID   = ""        # ← chat_id گیرنده (شناسه عددی)
TGRAM_TIMEOUT   = 10

# کلیدهای داخلی config (تغییر ندهید)
TGRAM_TOKEN_KEY  = "telegram_bot_token"
TGRAM_CHATID_KEY = "telegram_chat_id"

# ======================== CONFIG ========================

def load_config(config_path: Path) -> Dict[str, Any]:
    if not config_path.exists():
        return {}
    try:
        with open(config_path, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return {}

def save_config(config_path: Path, cfg: Dict[str, Any]) -> None:
    try:
        with open(config_path, "w", encoding="utf-8") as f:
            json.dump(cfg, f, ensure_ascii=False, indent=2)
    except Exception as e:
        logging.debug(f"Could not save config: {e}")


# ======================== GOOGLE DRIVE ========================

try:
    from google.oauth2.credentials import Credentials as _GCreds
    from google.auth.transport.requests import Request as _GRequest
    from google_auth_oauthlib.flow import InstalledAppFlow as _GFlow
    from googleapiclient.discovery import build as _gbuild
    from googleapiclient.http import MediaFileUpload as _MFU
    GDRIVE_AVAILABLE = True
except ImportError:
    GDRIVE_AVAILABLE = False
    logging.info("google-api-python-client نصب نیست — آپلود Drive غیرفعال است.")


def _drive_get_or_create_folder(service, name: str, parent_id: Optional[str] = None) -> str:
    """یک فولدر با نام داده‌شده پیدا یا می‌سازد و ID آن را برمی‌گرداند."""
    q = f"mimeType='application/vnd.google-apps.folder' and name={json.dumps(name)} and trashed=false"
    if parent_id:
        q += f" and '{parent_id}' in parents"
    res = service.files().list(q=q, fields="files(id)", pageSize=1).execute()
    files = res.get("files", [])
    if files:
        return files[0]["id"]
    meta = {"name": name, "mimeType": "application/vnd.google-apps.folder"}
    if parent_id:
        meta["parents"] = [parent_id]
    f = service.files().create(body=meta, fields="id").execute()
    return f["id"]


def _drive_upload_file(service, local_path: Path, parent_id: str) -> str:
    """فایل را آپلود می‌کند و لینک web viewable برمی‌گرداند."""
    import mimetypes
    mime, _ = mimetypes.guess_type(str(local_path))
    mime = mime or "application/octet-stream"
    meta = {"name": local_path.name, "parents": [parent_id]}
    media = _MFU(str(local_path), mimetype=mime, resumable=True)
    f = service.files().create(body=meta, media_body=media, fields="id,webViewLink").execute()
    return f.get("webViewLink", "")


def drive_build_service(token_path: Path) -> Any:
    """
    سرویس Drive می‌سازد. اگر token_path وجود داشت، از آن استفاده می‌کند.
    وگرنه OAuth flow را اجرا می‌کند (مرورگر باز می‌شود).
    token_path: مسیر ذخیره توکن refresh (مثلاً app_dir/drive_token.json)
    """
    if not GDRIVE_AVAILABLE:
        raise RuntimeError("google-api-python-client نصب نیست.")
    creds = None
    if token_path.exists():
        try:
            creds = _GCreds.from_authorized_user_file(str(token_path), DRIVE_SCOPES)
        except Exception:
            creds = None
    if creds and creds.expired and creds.refresh_token:
        try:
            creds.refresh(_GRequest())
        except Exception:
            creds = None
    if not creds or not creds.valid:
        if not OAUTH_CLIENT_ID or not OAUTH_CLIENT_SECRET:
            raise RuntimeError(
                "OAuth credentials در کد تعریف نشده‌اند.\n"
                "توسعه‌دهنده باید OAUTH_CLIENT_ID و OAUTH_CLIENT_SECRET را در ثابت‌ها تنظیم کند."
            )
        client_config = {
            "installed": {
                "client_id": OAUTH_CLIENT_ID,
                "client_secret": OAUTH_CLIENT_SECRET,
                "redirect_uris": [OAUTH_REDIRECT_URI],
                "auth_uri": "https://accounts.google.com/o/oauth2/auth",
                "token_uri": "https://oauth2.googleapis.com/token",
            }
        }
        flow = _GFlow.from_client_config(client_config, DRIVE_SCOPES)
        creds = flow.run_local_server(port=0)
        # ذخیره توکن برای استفاده‌های بعدی (بدون نیاز به لاگین مجدد)
        with open(token_path, "w") as f:
            f.write(creds.to_json())
    return _gbuild("drive", "v3", credentials=creds)


def drive_upload_run_outputs(
    service,
    out_dir: Path,
    input_pdfs: List[Path],
    subjects: List[str],
    ts: str,
    drive_root_name: str,
    log: Callable[[str], None],
) -> Dict[str, str]:
    """
    ساختار فولدر در Drive:
      <drive_root_name>/
        <موضوع‌ها>/
          <ts>/
            ورودی/   ← PDFهای اصلی
            خروجی/   ← Excel، Word، فولدرهای عکس

    خروجی: dict با کلیدهای "input_folder_link" و "output_folder_link"
    """
    subj_label = "، ".join(subjects) if subjects else "عمومی"

    # ساخت سلسله فولدرها
    root_id  = _drive_get_or_create_folder(service, drive_root_name)
    subj_id  = _drive_get_or_create_folder(service, subj_label,  root_id)
    ts_id    = _drive_get_or_create_folder(service, ts,           subj_id)
    in_id    = _drive_get_or_create_folder(service, "ورودی",      ts_id)
    out_id   = _drive_get_or_create_folder(service, "خروجی",      ts_id)

    links: Dict[str, str] = {}

    # ── آپلود PDFهای ورودی ──
    for pdf in input_pdfs:
        try:
            lnk = _drive_upload_file(service, pdf, in_id)
            log(f"     ☁ آپلود ورودی: {pdf.name}")
        except Exception as e:
            log(f"     ⚠ آپلود ورودی ناموفق ({pdf.name}): {str(e)[:60]}")

    # ── آپلود خروجی‌ها ──
    # ساختار: out_dir حاوی Excel/Word + زیرفولدرهای عکس
    def _upload_dir_recursive(local_dir: Path, parent_id: str):
        for item in sorted(local_dir.iterdir()):
            if item.is_file():
                try:
                    _drive_upload_file(service, item, parent_id)
                    log(f"     ☁ آپلود خروجی: {item.name}")
                except Exception as e:
                    log(f"     ⚠ آپلود خروجی ناموفق ({item.name}): {str(e)[:60]}")
            elif item.is_dir():
                sub_id = _drive_get_or_create_folder(service, item.name, parent_id)
                _upload_dir_recursive(item, sub_id)

    _upload_dir_recursive(out_dir, out_id)

    # لینک فولدرهای اصلی
    def _folder_link(fid: str) -> str:
        return f"https://drive.google.com/drive/folders/{fid}"

    links["input_folder_link"]  = _folder_link(in_id)
    links["output_folder_link"] = _folder_link(out_id)
    links["run_folder_link"]    = _folder_link(ts_id)
    return links


# ======================== TELEGRAM ========================

def telegram_send(token: str, chat_id: str, text: str) -> bool:
    """پیام متنی به تلگرام ارسال می‌کند. True=موفق، False=ناموفق."""
    if not token or not chat_id:
        return False
    try:
        url = f"https://api.telegram.org/bot{token}/sendMessage"
        resp = requests.post(
            url,
            json={"chat_id": chat_id, "text": text, "parse_mode": "HTML"},
            timeout=TGRAM_TIMEOUT,
        )
        return resp.status_code == 200
    except Exception as e:
        logging.debug(f"Telegram send failed: {e}")
        return False


def _tgram_notify(
    cfg: Dict[str, Any],
    user_label: str,
    subjects: List[str],
    ts: str,
    success: bool,
    drive_links: Optional[Dict[str, str]] = None,
    error_msg: str = "",
):
    """پیام اطلاع‌رسانی به تلگرام — توکن و Chat ID از ثابت‌های developer خوانده می‌شود."""
    # اول ثابت‌های developer را چک کن؛ اگر خالی بود config را بخوان (fallback)
    token   = TGRAM_BOT_TOKEN.strip() or cfg.get(TGRAM_TOKEN_KEY, "").strip()
    chat_id = TGRAM_CHAT_ID.strip()   or cfg.get(TGRAM_CHATID_KEY, "").strip()
    if not token or not chat_id:
        return

    subj = "، ".join(subjects) if subjects else "—"
    if success and drive_links:
        out_link = drive_links.get("output_folder_link", "")
        in_link  = drive_links.get("input_folder_link", "")
        msg = (
            f"✅ <b>اجرای موفق</b>\n"
            f"👤 کاربر: <code>{user_label}</code>\n"
            f"📂 موضوع: {subj}\n"
            f"🕐 زمان: {ts}\n\n"
            f"📥 <a href='{in_link}'>فولدر ورودی</a>\n"
            f"📤 <a href='{out_link}'>فولدر خروجی</a>"
        )
    elif success:
        msg = (
            f"✅ <b>اجرای موفق</b> (آپلود Drive غیرفعال)\n"
            f"👤 کاربر: <code>{user_label}</code>\n"
            f"📂 موضوع: {subj}\n"
            f"🕐 زمان: {ts}"
        )
    else:
        msg = (
            f"❌ <b>آپلود Drive ناموفق</b>\n"
            f"👤 کاربر: <code>{user_label}</code>\n"
            f"📂 موضوع: {subj}\n"
            f"🕐 زمان: {ts}\n"
            f"⚠ خطا: {error_msg[:200]}"
        )
    telegram_send(token, chat_id, msg)


def load_subjects(cfg: Dict[str, Any], timeout: int = SUBJECTS_FETCH_TIMEOUT) -> List[str]:
    """
    لیست موضوع‌ها را از SUBJECTS_URL (raw گیت‌هاب) می‌خواند.
    در صورت خطا/آفلاین به آخرین کشِ موفق داخل config برمی‌گردد.
    خروجی موفق در cfg["_subjects_cache"] ذخیره می‌شود (caller باید config را save کند).
    """
    try:
        import requests as _rq
        r = _rq.get(SUBJECTS_URL, timeout=timeout)
        r.raise_for_status()
        data = r.json()
        subs = data.get("subjects", []) if isinstance(data, dict) else data
        subs = [str(s).strip() for s in subs if str(s).strip()]
        if subs:
            cfg["_subjects_cache"] = subs
            logging.info(f"Subjects loaded from remote: {len(subs)}")
            return subs
    except Exception as e:
        logging.info(f"Subjects remote fetch failed, using cache: {e}")
    return list(cfg.get("_subjects_cache", []))

# ======================== API KEY MANAGER ========================

class APIKeyManager:
    def __init__(self, api_keys: List[str], db_path: Optional[Path] = None):
        self.all_keys = api_keys.copy()
        self.valid_keys: deque = deque()
        self.invalid_keys: set = set()
        self.quota_exceeded_keys: set = set()
        self.lock = Lock()
        self.db_path = db_path

        if db_path:
            self._init_db()
            self._load_key_status()

        self._initialize_valid_keys()

    def _init_db(self) -> None:
        self.db_path.parent.mkdir(parents=True, exist_ok=True)
        with sqlite3.connect(str(self.db_path)) as conn:
            conn.execute("""
                CREATE TABLE IF NOT EXISTS api_key_status (
                    key_hash TEXT PRIMARY KEY,
                    is_valid INTEGER,
                    is_quota_exceeded INTEGER,
                    last_error TEXT,
                    last_checked REAL
                )
            """)
            conn.commit()

    def _load_key_status(self) -> None:
        try:
            with sqlite3.connect(str(self.db_path)) as conn:
                cur = conn.execute(
                    "SELECT key_hash, is_valid, is_quota_exceeded FROM api_key_status"
                )
                for key_hash, is_valid, is_quota_exceeded in cur.fetchall():
                    if not is_valid:
                        self.invalid_keys.add(key_hash)
                    if is_quota_exceeded:
                        self.quota_exceeded_keys.add(key_hash)
        except Exception as e:
            logging.debug(f"Could not load key status: {e}")

    def _save_key_status(self, key_hash: str, is_valid: bool, is_quota: bool, error: str = "") -> None:
        if not self.db_path:
            return
        try:
            with sqlite3.connect(str(self.db_path)) as conn:
                conn.execute("""
                    INSERT INTO api_key_status (key_hash, is_valid, is_quota_exceeded, last_error, last_checked)
                    VALUES (?, ?, ?, ?, ?)
                    ON CONFLICT(key_hash) DO UPDATE SET
                        is_valid=excluded.is_valid,
                        is_quota_exceeded=excluded.is_quota_exceeded,
                        last_error=excluded.last_error,
                        last_checked=excluded.last_checked
                """, (key_hash, int(is_valid), int(is_quota), error, time.time()))
                conn.commit()
        except Exception as e:
            logging.debug(f"Could not save key status: {e}")

    def _hash(self, key: str) -> str:
        return hashlib.md5(key.encode()).hexdigest()

    def _initialize_valid_keys(self) -> None:
        for key in self.all_keys:
            h = self._hash(key)
            if h not in self.invalid_keys and h not in self.quota_exceeded_keys:
                self.valid_keys.append(key)
        logging.info(f"API Key Manager: {len(self.valid_keys)}/{len(self.all_keys)} keys active")

    def get_next_key(self) -> Optional[str]:
        with self.lock:
            return self.valid_keys[0] if self.valid_keys else None

    def rotate_key(self) -> None:
        with self.lock:
            if self.valid_keys:
                key = self.valid_keys.popleft()
                self.valid_keys.append(key)

    def mark_invalid(self, key: str, error: str = "") -> None:
        h = self._hash(key)
        with self.lock:
            if key in self.valid_keys:
                self.valid_keys.remove(key)
            self.invalid_keys.add(h)
        self._save_key_status(h, False, False, error)
        logging.warning(f"Key marked invalid: {error[:60]}")

    def mark_quota_exceeded(self, key: str) -> None:
        h = self._hash(key)
        with self.lock:
            if key in self.valid_keys:
                self.valid_keys.remove(key)
            self.quota_exceeded_keys.add(h)
        self._save_key_status(h, True, True, "Quota exceeded")
        logging.warning("Key marked quota-exceeded, trying next key")

    def mark_success(self, key: str) -> None:
        h = self._hash(key)
        with self.lock:
            if key in self.valid_keys:
                self.valid_keys.remove(key)
            self.valid_keys.appendleft(key)
            self.quota_exceeded_keys.discard(h)
        self._save_key_status(h, True, False, "")

    @property
    def active_count(self) -> int:
        return len(self.valid_keys)

# ======================== GEMINI API ========================

def pdf_to_base64(pdf_path: Path) -> str:
    with open(pdf_path, "rb") as f:
        return base64.b64encode(f.read()).decode("utf-8")


def build_merg_sample_base64(pdf_path: Path, page_cap: int = MERG_SAMPLE_PAGE_CAP) -> tuple:
    if not PYPDF_AVAILABLE:
        return pdf_to_base64(pdf_path), -1, []

    try:
        reader = PdfReader(str(pdf_path))
        total = len(reader.pages)
    except Exception as e:
        logging.warning(f"خواندن MERG.pdf برای نمونه‌برداری ناموفق بود: {e}")
        return pdf_to_base64(pdf_path), -1, []

    if total <= page_cap:
        return pdf_to_base64(pdf_path), total, list(range(1, total + 1))

    step = (total - 1) / (page_cap - 1)
    indices = sorted(set(round(i * step) for i in range(page_cap)))
    indices = [max(0, min(total - 1, idx)) for idx in indices]
    indices = sorted(set(indices))

    writer = PdfWriter()
    for idx in indices:
        writer.add_page(reader.pages[idx])

    import io
    buf = io.BytesIO()
    writer.write(buf)
    b64 = base64.b64encode(buf.getvalue()).decode("utf-8")
    return b64, total, [i + 1 for i in indices]


def _call_gemini_core(
    parts: List[Dict[str, Any]],
    key_manager: APIKeyManager,
    model: str = "gemini-2.0-flash",
    last_call_time: Optional[List[float]] = None,
    cancel_event: Optional[Event] = None,
    parse_json: bool = True,
) -> Any:
    """هسته‌ی فراخوانی Gemini با هر ترکیبی از parts (PDF / عکس / متن)."""
    url = f"https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent"

    def _sleep(seconds: float):
        if cancel_event is None:
            time.sleep(seconds)
            return
        end = time.time() + seconds
        while time.time() < end:
            if cancel_event.is_set():
                return
            time.sleep(min(0.1, end - time.time()))

    body = {
        "contents": [
            {
                "role": "user",
                "parts": parts,
            }
        ],
        "generationConfig": {
            "temperature": TEMPERATURE,
            "maxOutputTokens": 65536,
        },
        "safetySettings": [
            {"category": "HARM_CATEGORY_HARASSMENT", "threshold": "BLOCK_NONE"},
            {"category": "HARM_CATEGORY_HATE_SPEECH", "threshold": "BLOCK_NONE"},
            {"category": "HARM_CATEGORY_SEXUALLY_EXPLICIT", "threshold": "BLOCK_NONE"},
            {"category": "HARM_CATEGORY_DANGEROUS_CONTENT", "threshold": "BLOCK_NONE"},
        ],
    }

    if last_call_time is None:
        last_call_time = [0.0]

    keys_tried = 0
    max_keys = len(key_manager.all_keys)
    all_errors: List[str] = []

    while keys_tried < max_keys:
        if cancel_event is not None and cancel_event.is_set():
            raise RuntimeError("Cancelled")

        api_key = key_manager.get_next_key()
        if not api_key:
            break

        keys_tried += 1
        key_num = key_manager.all_keys.index(api_key) + 1
        headers = {"Content-Type": "application/json", "x-goog-api-key": api_key}

        for attempt in range(MAX_RETRIES_PER_KEY):
            if cancel_event is not None and cancel_event.is_set():
                raise RuntimeError("Cancelled")

            now = time.time()
            wait = API_RATE_DELAY - (now - last_call_time[0])
            if wait > 0:
                _sleep(wait)
            last_call_time[0] = time.time()

            try:
                logging.info(f"API: key {key_num}/{max_keys}, attempt {attempt+1}/{MAX_RETRIES_PER_KEY}")
                resp = requests.post(url, headers=headers, json=body, timeout=TIMEOUT_S)

                if resp.status_code == 429:
                    logging.info(f"Key {key_num}: quota exceeded → next key")
                    key_manager.mark_quota_exceeded(api_key)
                    all_errors.append(f"Key {key_num}: 429 Quota")
                    break

                if resp.status_code in [401, 403]:
                    logging.warning(f"Key {key_num}: auth error {resp.status_code} → invalid")
                    key_manager.mark_invalid(api_key, f"HTTP {resp.status_code}")
                    all_errors.append(f"Key {key_num}: {resp.status_code} Auth")
                    break

                resp.raise_for_status()
                data = resp.json()

                candidates = data.get("candidates", [])
                if not candidates:
                    raise ValueError("No candidates in response")

                finish_reason = candidates[0].get("finishReason", "STOP")
                if finish_reason not in ("STOP", "MAX_TOKENS"):
                    raise ValueError(f"Bad finishReason: {finish_reason}")

                parts = candidates[0].get("content", {}).get("parts", [])
                if not parts or "text" not in parts[0]:
                    raise ValueError("No text in response")

                text = parts[0]["text"].strip()

                if not parse_json:
                    key_manager.mark_success(api_key)
                    logging.info(f"Key {key_num}: success ✓ (text)")
                    return text

                text = re.sub(r"```json\s*", "", text)
                text = re.sub(r"```\s*", "", text)

                m = re.search(r"[\[\{][\s\S]*[\]\}]", text)
                if not m:
                    raise ValueError(f"No JSON in response: {text[:200]}")

                result = json.loads(m.group())
                key_manager.mark_success(api_key)
                logging.info(f"Key {key_num}: success ✓")
                return result

            except requests.Timeout:
                err = f"Key {key_num} attempt {attempt+1}: Timeout"
                all_errors.append(err)
                if attempt < MAX_RETRIES_PER_KEY - 1:
                    logging.info(f"{err}, retrying...")
                    _sleep(RETRY_DELAY * (attempt + 1))
                else:
                    logging.warning(f"{err}, next key")
                    key_manager.rotate_key()

            except requests.ConnectionError:
                err = f"Key {key_num} attempt {attempt+1}: ConnectionError"
                all_errors.append(err)
                if attempt < MAX_RETRIES_PER_KEY - 1:
                    logging.info(f"{err}, retrying...")
                    _sleep(RETRY_DELAY * (attempt + 1))
                else:
                    logging.warning(f"{err}, next key")
                    key_manager.rotate_key()

            except json.JSONDecodeError:
                err = f"Key {key_num} attempt {attempt+1}: JSON parse error"
                all_errors.append(err)
                if attempt < MAX_RETRIES_PER_KEY - 1:
                    _sleep(RETRY_DELAY)
                else:
                    key_manager.rotate_key()
                    break

            except Exception as e:
                err = f"Key {key_num} attempt {attempt+1}: {str(e)[:100]}"
                all_errors.append(err)
                if "401" in str(e) or "403" in str(e):
                    key_manager.mark_invalid(api_key, str(e)[:100])
                    break
                if attempt < MAX_RETRIES_PER_KEY - 1:
                    _sleep(RETRY_DELAY * (attempt + 1))
                else:
                    key_manager.rotate_key()
                    break

            if cancel_event is not None and cancel_event.is_set():
                raise RuntimeError("Cancelled")

    raise RuntimeError(
        f"All {keys_tried} API keys failed. "
        f"Errors: {all_errors[-3:] if all_errors else []}"
    )


def call_gemini_with_pdf(
    pdf_base64: str,
    prompt: str,
    key_manager: APIKeyManager,
    model: str = "gemini-2.0-flash",
    last_call_time: Optional[List[float]] = None,
    cancel_event: Optional[Event] = None,
    parse_json: bool = True,
) -> Any:
    parts = [
        {"inline_data": {"mime_type": "application/pdf", "data": pdf_base64}},
        {"text": prompt},
    ]
    return _call_gemini_core(parts, key_manager, model, last_call_time, cancel_event, parse_json)


def call_gemini_with_images(
    images: List[tuple],          # [(base64, mime_type), ...]
    prompt: str,
    key_manager: APIKeyManager,
    model: str = "gemini-2.0-flash",
    last_call_time: Optional[List[float]] = None,
    cancel_event: Optional[Event] = None,
    parse_json: bool = True,
) -> Any:
    parts: List[Dict[str, Any]] = [
        {"inline_data": {"mime_type": mime, "data": b64}} for (b64, mime) in images
    ]
    parts.append({"text": prompt})
    return _call_gemini_core(parts, key_manager, model, last_call_time, cancel_event, parse_json)


# ======================== PROMPTS ========================

def build_feature_discovery_prompt(category: str = "", is_sample: bool = False) -> str:
    cat_line = f'\nنوع کالا/دسته مورد نظر: {category.strip()}\n' if category.strip() else ""
    if is_sample:
        intro = (
            "این PDF حاصل چسباندن چند کاتالوگ محصول مختلف به‌هم است و چند صفحه‌ی "
            "نمونه از نقاط مختلف آن استخراج شده (نه کل فایل).\n"
            "مهم: همه‌ی این صفحات نمونه را بررسی کن — ممکن است هر صفحه به یک محصول متفاوت تعلق داشته باشد."
        )
    else:
        intro = "این PDF یک کاتالوگ محصول است.\n\nمهم: تمام صفحات PDF را از ابتدا تا انتها بخوان."
    return f"""{intro}
{cat_line}
وظیفه: مهم‌ترین ویژگی‌های فنی/توصیفی مشترکی که برای محصولات این کاتالوگ تکرار می‌شوند را پیدا کن.
این ویژگی‌ها بعداً به‌عنوان ستون‌های جدول استخراج استفاده می‌شوند.

خروجی فقط یک آرایه JSON از رشته‌ها باشد — بدون هیچ متن اضافه:
مثال:
["جنس", "سایز", "فشار کاری", "رنگ", "وزن"]

JSON:"""


def build_extraction_prompt(features: List[str], category: str = "") -> str:
    features_str = "\n".join(f"- {f}" for f in features)
    cat_line = f'\nنوع کالا/دسته مورد نظر: {category.strip()}\n' if category.strip() else ""
    return f"""این PDF یک کاتالوگ محصول است.

مهم: تمام صفحات PDF را از ابتدا تا انتها بخوان و بررسی کن.
{cat_line}
تمام آیتم‌های محصول را پیدا کن (هر کد، مدل، سایز یا variant جداگانه)
و برای هر آیتم ویژگی‌های زیر را استخراج کن:
{features_str}

قوانین خروجی:
1) یک آرایه JSON از objects — بدون هیچ متن اضافه‌ای قبل یا بعد
2) کلید اول هر object دقیقاً "Model" باشد (نام/کد آیتم طبق PDF)
3) بقیه کلیدها دقیقاً همان متن ویژگی‌های بالا باشند
4) اگر مقداری پیدا نشد: null
5) اگر یک آیتم چند variant دارد، هر کدام ردیف جداگانه

مثال:
[
  {{"Model": "کد آیتم", "{features[0] if features else 'ویژگی'}": "مقدار"}},
  {{"Model": "کد آیتم ۲", "{features[0] if features else 'ویژگی'}": null}}
]

JSON:"""


def build_model_search_prompt(models: List[str], features: List[str], category: str = "") -> str:
    features_str = "\n".join(f"- {f}" for f in features)
    cat_line = f'\nنوع کالا/دسته/موضوع مورد نظر: {category.strip()}\n' if category.strip() else ""

    if models:
        models_str = "\n".join(f"- {m}" for m in models)
        models_block = f"""مدل‌هایی که باید پیدا کنی:
{models_str}

قانون اضافه: فقط مدل‌هایی که در لیست بالا هستند را برگردان."""
    else:
        models_block = "هیچ لیست مدل خاصی مشخص نشده — تمام مدل‌ها/آیتم‌های موجود در PDF را پیدا کن."

    return f"""این PDF یک کاتالوگ محصول است.
{cat_line}
مهم: تمام صفحات PDF را از ابتدا تا انتها بخوان و بررسی کن.

{models_block}

ویژگی‌هایی که برای هر مدل باید استخراج شوند:
{features_str}

قوانین خروجی:
1) یک آرایه JSON باشد — بدون هیچ متن اضافه‌ای قبل یا بعد
2) هر آبجکت یک ردیف = یک مدل یا یک variant از آن
3) کلید اول هر آبجکت دقیقاً "Model" باشد (نام دقیق مدل طبق PDF)
4) بقیه کلیدها دقیقاً همان متن ویژگی‌های بالا باشند
5) اگر اطلاعاتی پیدا نشد: null
6) اگر یک مدل چند سری/variant دارد، هر کدام ردیف جداگانه

مثال:
[
  {{"Model": "XR-100", "{features[0] if features else 'ویژگی ۱'}": "مقدار"}},
  {{"Model": "XR-200A", "{features[0] if features else 'ویژگی ۱'}": null}}
]

JSON:"""


# ======================== CORE EXTRACTION ========================

def discover_features(
    merg_pdf: Path,
    category: str,
    key_manager: APIKeyManager,
    model: str,
    last_call_time: List[float],
    cancel_event: Optional[Event] = None,
    log_callback: Optional[Callable[[str], None]] = None,
) -> List[str]:
    b64, total_pages, sampled = build_merg_sample_base64(merg_pdf)
    is_sample = total_pages > 0 and len(sampled) < total_pages
    if log_callback:
        if is_sample:
            log_callback(f"     (نمونه‌برداری: {len(sampled)} از {total_pages} صفحه)")
        elif total_pages > 0:
            log_callback(f"     ({total_pages} صفحه — کل فایل ارسال می‌شود)")

    prompt = build_feature_discovery_prompt(category, is_sample=is_sample)
    raw = call_gemini_with_pdf(b64, prompt, key_manager, model, last_call_time, cancel_event)
    if isinstance(raw, list):
        return [str(x).strip() for x in raw if str(x).strip()]
    if isinstance(raw, dict):
        return [k for k in raw.keys() if k != "Model"]
    return []


def extract_pdf(
    pdf_path: Path,
    features: List[str],
    category: str,
    key_manager: APIKeyManager,
    model: str,
    last_call_time: List[float],
    cancel_event: Optional[Event] = None,
) -> List[Dict[str, Any]]:
    b64 = pdf_to_base64(pdf_path)
    prompt = build_extraction_prompt(features, category)
    raw = call_gemini_with_pdf(b64, prompt, key_manager, model, last_call_time, cancel_event)
    rows: List[Dict[str, Any]] = []
    if isinstance(raw, list):
        rows = [r for r in raw if isinstance(r, dict)]
    elif isinstance(raw, dict):
        rows = [raw]
    for r in rows:
        r.setdefault("_source", pdf_path.name)
    return rows


def extract_pdf_model_mode(
    pdf_path: Path,
    models: List[str],
    features: List[str],
    key_manager: APIKeyManager,
    model: str,
    last_call_time: List[float],
    cancel_event: Optional[Event] = None,
    category: str = "",
) -> List[Dict[str, Any]]:
    b64 = pdf_to_base64(pdf_path)
    prompt = build_model_search_prompt(models, features, category)
    raw = call_gemini_with_pdf(b64, prompt, key_manager, model, last_call_time, cancel_event)
    rows: List[Dict[str, Any]] = []
    if isinstance(raw, list):
        rows = [r for r in raw if isinstance(r, dict)]
    elif isinstance(raw, dict):
        rows = [raw]
    for r in rows:
        r.setdefault("_source", pdf_path.name)
    return rows


# ======================== OUTPUT PATH MANAGEMENT ========================

def make_pdf_out_dir(pdf_path: Path) -> Path:
    """
    فولدر خروجی مربوط به یک PDF می‌سازد و برمی‌گرداند.
    نام: <pdf_stem>_out  در کنار همان PDF.
    مثال: /فولدر/محصول_A.pdf  →  /فولدر/محصول_A_out/
    """
    out = pdf_path.parent / f"{pdf_path.stem}_out"
    out.mkdir(parents=True, exist_ok=True)
    return out


def make_folder_out_dir(folder_path: Path, ts: str) -> Path:
    """
    فولدر خروجی برای حالت فولدر (شامل چند PDF).
    نام: <folder_name>_out  در کنار همان فولدر.
    """
    out = folder_path.parent / f"{folder_path.name}_out"
    out.mkdir(parents=True, exist_ok=True)
    return out


# ======================== INPUT DISCOVERY ========================

def is_merg(p: Path) -> bool:
    return p.name.lower() == MERG_FILENAME.lower()


def collect_folder_pdfs(folder: Path) -> List[Path]:
    return sorted(
        p for p in folder.glob("*.pdf")
        if not is_merg(p)
    )


def find_merg(folder: Path) -> Optional[Path]:
    for p in folder.glob("*.pdf"):
        if is_merg(p):
            return p
    return None


# ======================== EXCEL OUTPUT ========================

def _style_sheet(ws, all_cols: List[str], rows: List[Dict[str, Any]]) -> None:
    header_font = Font(name="Arial", bold=True, color="FFFFFF", size=11)
    header_fill = PatternFill("solid", start_color="063f47")
    header_align = Alignment(horizontal="center", vertical="center", wrap_text=True)
    data_font = Font(name="Arial", size=10)
    data_align = Alignment(horizontal="right", vertical="center", wrap_text=True)
    alt_fill = PatternFill("solid", start_color="E8F4F6")
    thin = Side(style="thin", color="CCCCCC")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    for col_idx, col_name in enumerate(all_cols, 1):
        cell = ws.cell(row=1, column=col_idx, value=col_name)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = header_align
        cell.border = border
    ws.row_dimensions[1].height = 35

    for row_idx, row_data in enumerate(rows, 2):
        use_alt = row_idx % 2 == 0
        for col_idx, col_name in enumerate(all_cols, 1):
            val = row_data.get(col_name, "")
            if val is None:
                val = ""
            cell = ws.cell(row=row_idx, column=col_idx, value=str(val) if val != "" else "")
            cell.font = data_font
            cell.alignment = data_align
            cell.border = border
            if use_alt:
                cell.fill = alt_fill
        ws.row_dimensions[row_idx].height = 20

    ws.column_dimensions["A"].width = 22
    for col_idx in range(2, len(all_cols) + 1):
        letter = ws.cell(row=1, column=col_idx).column_letter
        ws.column_dimensions[letter].width = 18
    ws.freeze_panes = "A2"


def create_excel_output(
    results: List[Dict[str, Any]],
    features: List[str],
    output_path: Path,
    include_source: bool = True,
) -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = "نتایج"

    all_cols = ["Model"] + features
    if include_source:
        all_cols = all_cols + ["فایل منبع"]
    rows = []
    for r in results:
        rr = dict(r)
        if include_source:
            rr["فایل منبع"] = r.get("_source", "")
        rows.append(rr)

    _style_sheet(ws, all_cols, rows)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(str(output_path))


# ======================== TEXT → WORD + IMAGE EXTRACTION ========================

WORD_FONT = "Tahoma"            # فونت سازگار با فارسی و موجود روی ویندوز
IMAGE_MIN_BYTES = 100           # فقط اسپیسرهای بسیار کوچک (۱px) نادیده گرفته می‌شوند

_IMG_EXTS = (".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tiff", ".webp")


def _mime_for_ext(ext: str) -> str:
    return {
        ".png": "image/png", ".jpg": "image/jpeg", ".jpeg": "image/jpeg",
        ".gif": "image/gif", ".bmp": "image/bmp", ".webp": "image/webp",
        ".tiff": "image/tiff",
    }.get(ext.lower(), "image/png")


def build_text_extraction_prompt() -> str:
    return """این یک فایل PDF (کاتالوگ/سند) است.

وظیفه: کل محتوای متنی PDF را از ابتدا تا انتها، به‌ترتیب و کامل استخراج کن و
به‌صورت تمیز، ساختاریافته و پاراگراف‌بندی‌شده در قالب Markdown بازنویسی کن.

قوانین:
- تمام صفحات را به ترتیب بخوان؛ چیزی را جا نینداز.
- در ابتدای محتوای هر صفحه، دقیقاً یک نشانگر در یک خط جداگانه بگذار به این شکل:
  <!--page:شماره_صفحه-->
  (مثلاً برای صفحه ۳ بنویس: <!--page:3--> ). این نشانگر فقط برای موقعیت‌یابی است.
- ساختار سند را حفظ کن: عنوان بخش‌ها را با #، ## و ### مشخص کن.
- متن‌های پیوسته را به پاراگراف‌های منظم و خوانا تقسیم کن.
- جدول‌ها (مثل جدول مشخصات فنی) را به‌صورت جدول Markdown بنویس (با | و سطر جداکننده ---).
- فهرست‌ها را با «- » یا شماره‌گذاری بنویس.
- هدر/فوتر تکراری صفحات و شماره‌ی صفحه را حذف کن (نشانگر page را نگه‌دار).
- هیچ توضیح یا متنی از خودت اضافه نکن؛ فقط محتوای خودِ سند.
- خروجی فقط Markdown خام باشد، بدون ``` و بدون هیچ مقدمه‌ای.

محتوا:"""


def extract_pdf_text_via_gemini(
    pdf_path: Path,
    key_manager: APIKeyManager,
    model: str,
    last_call_time: List[float],
    cancel_event: Optional[Event] = None,
) -> str:
    """یک API call: متن تمیزِ Markdown از کل PDF برمی‌گرداند."""
    b64 = pdf_to_base64(pdf_path)
    prompt = build_text_extraction_prompt()
    raw = call_gemini_with_pdf(
        b64, prompt, key_manager, model, last_call_time, cancel_event, parse_json=False
    )
    text = (raw or "").strip()
    text = re.sub(r"^```(?:markdown|md)?\s*\n?", "", text)
    text = re.sub(r"\n?```\s*$", "", text)
    return text.strip()


# ---- توضیح‌گذاری تصاویر با Gemini Vision (دسته‌ای) ----

def build_image_desc_prompt(count: int, subjects: List[str]) -> str:
    types_str = "، ".join(IMAGE_TYPES)
    subj = "، ".join([s for s in (subjects or []) if s])
    subj_line = f"\nموضوع/دسته‌ی این کاتالوگ: {subj}\n" if subj else ""
    return f"""به تو {count} تصویرِ استخراج‌شده از یک کاتالوگ محصول داده می‌شود، دقیقاً به همان ترتیب.
{subj_line}
برای هر تصویر به‌ترتیب یک object بساز با این کلیدها:
- "type": یکی از این مقادیر: {types_str}
- "product": نام یا کد کالای مرتبط با تصویر اگر قابل تشخیص است، در غیر این صورت null
- "caption": یک توضیح کوتاه فارسی (حداکثر یک جمله) درباره‌ی محتوای تصویر

قوانین خروجی:
1) فقط یک آرایه JSON برگردان — بدون هیچ متن اضافه‌ای قبل یا بعد
2) دقیقاً {count} object، به همان ترتیب تصاویر
3) اگر چیزی قابل تشخیص نبود، مقدار null بگذار

JSON:"""


def describe_images_via_gemini(
    unique_records: List[Dict[str, Any]],
    img_dir: Path,
    key_manager: APIKeyManager,
    model: str,
    last_call_time: List[float],
    cancel_event: Optional[Event],
    subjects: List[str],
) -> Dict[str, Dict[str, Any]]:
    """دسته‌ای: برای هر تصویرِ یکتا {type, product, caption} برمی‌گرداند (کلید = filename)."""
    desc_map: Dict[str, Dict[str, Any]] = {}
    for start in range(0, len(unique_records), IMAGE_DESC_BATCH):
        if cancel_event is not None and cancel_event.is_set():
            break
        batch = unique_records[start:start + IMAGE_DESC_BATCH]
        images: List[tuple] = []
        for rec in batch:
            try:
                with open(img_dir / rec["filename"], "rb") as f:
                    data = f.read()
            except Exception:
                data = b""
            images.append((base64.b64encode(data).decode("utf-8"), _mime_for_ext(rec["ext"])))

        prompt = build_image_desc_prompt(len(batch), subjects)
        try:
            raw = call_gemini_with_images(
                images, prompt, key_manager, model, last_call_time, cancel_event, parse_json=True
            )
        except Exception as e:
            logging.warning(f"image-desc batch failed: {str(e)[:80]}")
            raw = []
        arr = raw if isinstance(raw, list) else []
        for i, rec in enumerate(batch):
            d = arr[i] if (i < len(arr) and isinstance(arr[i], dict)) else {}
            prod = d.get("product")
            desc_map[rec["filename"]] = {
                "type": (str(d.get("type") or "").strip() or "سایر"),
                "product": (str(prod).strip() if prod not in (None, "", "null") else None),
                "caption": str(d.get("caption") or "").strip(),
            }
    for rec in unique_records:
        desc_map.setdefault(rec["filename"], {"type": "سایر", "product": None, "caption": ""})
    return desc_map


def embed_image_metadata(path: Path, desc: Dict[str, Any]) -> None:
    """توضیحات را در متادیتای خود فایل عکس درج می‌کند (best-effort)."""
    if not PIL_AVAILABLE:
        return
    try:
        ext = path.suffix.lower()
        caption = desc.get("caption", "") or ""
        itype = desc.get("type", "") or ""
        product = desc.get("product") or ""
        summary = f"type={itype}; product={product}; caption={caption}"
        img = Image.open(path)
        if ext == ".png" and PngImagePlugin is not None:
            meta = PngImagePlugin.PngInfo()
            meta.add_text("Description", caption)
            meta.add_text("Type", itype)
            meta.add_text("Product", str(product))
            img.save(path, pnginfo=meta)
        elif ext in (".jpg", ".jpeg"):
            exif = img.getexif()
            # ImageDescription فقط مقدار ASCII امن (کد کالا) برای سازگاری حداکثری
            try:
                ascii_desc = f"type={itype}; product={product}".encode("ascii", "ignore").decode("ascii")
                exif[0x010E] = ascii_desc
            except Exception:
                pass
            # توضیح کامل فارسی در UserComment با پیشوند استاندارد UNICODE
            try:
                ifd = exif.get_ifd(0x8769)
                full = f"{itype} | کالا: {product} | {caption}".strip()
                ifd[0x9286] = b"UNICODE\x00" + full.encode("utf-16-le")
            except Exception:
                pass
            img.save(path, exif=exif)
        # سایر فرمت‌ها: رد می‌شود (sidecar همه را دارد)
    except Exception as e:
        logging.debug(f"metadata embed failed {path.name}: {e}")


# ---- استخراج تصاویر (نام یکتای ۳۲ کاراکتری) ----

def extract_images_from_pdf(pdf_path: Path, out_dir: Path) -> tuple:
    """
    تصاویر PDF را با نام یکتای ۳۲ کاراکتری ذخیره می‌کند.
    خروجی: (unique_records, occurrences)
      unique_records: [{uuid, filename, ext, bytes, hash, first_page}]
      occurrences:    [{page, filename}]  (هر بار ظاهر شدن، برای رفرنس درون‌متن)
    تصاویر تکراری (هش یکسان) یک‌بار ذخیره می‌شوند ولی همه‌ی occurrenceها ثبت می‌شوند.
    """
    if not PYPDF_AVAILABLE:
        return [], []
    try:
        reader = PdfReader(str(pdf_path))
    except Exception as e:
        logging.warning(f"خواندن PDF برای استخراج عکس ناموفق بود ({pdf_path.name}): {e}")
        return [], []

    out_dir.mkdir(parents=True, exist_ok=True)
    hash_to_rec: Dict[str, Dict[str, Any]] = {}
    unique_records: List[Dict[str, Any]] = []
    occurrences: List[Dict[str, Any]] = []

    for page_num, page in enumerate(reader.pages, 1):
        try:
            images = list(page.images)
        except Exception as e:
            logging.debug(f"page.images failed p{page_num} of {pdf_path.name}: {e}")
            continue

        for img in images:
            try:
                data = img.data
            except Exception:
                continue
            if not data or len(data) < IMAGE_MIN_BYTES:
                continue

            h = hashlib.md5(data).hexdigest()
            rec = hash_to_rec.get(h)
            if rec is None:
                name = getattr(img, "name", "") or ""
                ext = Path(name).suffix.lower()
                if ext not in _IMG_EXTS:
                    ext = ".png"
                uid = uuid.uuid4().hex            # ۳۲ کاراکتر یکتا
                filename = f"{uid}{ext}"
                try:
                    with open(out_dir / filename, "wb") as f:
                        f.write(data)
                except Exception as e:
                    logging.debug(f"could not write image {filename}: {e}")
                    continue
                rec = {
                    "uuid": uid, "filename": filename, "ext": ext,
                    "bytes": len(data), "hash": h, "first_page": page_num,
                }
                hash_to_rec[h] = rec
                unique_records.append(rec)

            occurrences.append({"page": page_num, "filename": rec["filename"]})

    return unique_records, occurrences


def write_images_index(
    img_dir: Path,
    pdf_name: str,
    unique_records: List[Dict[str, Any]],
    occurrences: List[Dict[str, Any]],
    desc_map: Dict[str, Dict[str, Any]],
) -> List[Dict[str, Any]]:
    pages_by_file: Dict[str, List[int]] = {}
    for occ in occurrences:
        pages_by_file.setdefault(occ["filename"], []).append(occ["page"])

    items: List[Dict[str, Any]] = []
    for rec in unique_records:
        d = desc_map.get(rec["filename"], {})
        items.append({
            "filename": rec["filename"],
            "type": d.get("type", ""),
            "product": d.get("product"),
            "caption": d.get("caption", ""),
            "pages": sorted(set(pages_by_file.get(rec["filename"], []))),
            "bytes": rec.get("bytes", 0),
        })
    try:
        with open(img_dir / "images_index.json", "w", encoding="utf-8") as f:
            json.dump({"pdf": pdf_name, "count": len(items), "images": items},
                      f, ensure_ascii=False, indent=2)
    except Exception as e:
        logging.debug(f"could not write images_index.json: {e}")
    return items


# ---- Markdown → DOCX (RTL) ----

def _set_paragraph_rtl(paragraph) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    pPr.append(bidi)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def _add_md_runs(paragraph, text: str) -> None:
    for part in re.split(r"(\*\*.+?\*\*)", text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def _add_hyperlink(paragraph, text: str, target: str) -> bool:
    """هایپرلینک نسبی (best-effort). در صورت خطا False برمی‌گرداند."""
    try:
        part = paragraph.part
        r_id = part.relate_to(
            target,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True,
        )
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)
        run = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        color = OxmlElement("w:color"); color.set(qn("w:val"), "0563C1"); rPr.append(color)
        u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rPr.append(u)
        run.append(rPr)
        t = OxmlElement("w:t"); t.text = text; run.append(t)
        hyperlink.append(run)
        paragraph._p.append(hyperlink)
        return True
    except Exception:
        return False


def _is_table_sep(line: str) -> bool:
    s = line.strip().strip("|")
    return bool(s) and all(c in " -:|" for c in s) and "-" in s


def _parse_table_row(line: str) -> List[str]:
    return [c.strip() for c in line.strip().strip("|").split("|")]


def _add_table(doc, header: List[str], rows: List[List[str]]) -> None:
    cols = len(header)
    if cols == 0:
        return
    table = doc.add_table(rows=1, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    table.style = "Table Grid"
    tblPr = table._tbl.tblPr
    tblPr.append(OxmlElement("w:bidiVisual"))

    hdr = table.rows[0].cells
    for i, htext in enumerate(header[:cols]):
        p = hdr[i].paragraphs[0]
        _set_paragraph_rtl(p)
        r = p.add_run(htext)
        r.bold = True

    for row in rows:
        cells = table.add_row().cells
        for i in range(cols):
            val = row[i] if i < len(row) else ""
            p = cells[i].paragraphs[0]
            _set_paragraph_rtl(p)
            _add_md_runs(p, val)


_PAGE_MARK_RE = re.compile(r"<!--\s*page\s*:\s*(\d+)\s*-->")


def _flush_page_images(doc, page_images, page, img_rel_dir) -> None:
    if not page_images or page is None:
        return
    refs = page_images.get(page)
    if not refs:
        return
    for filename, d in refs:
        p = doc.add_paragraph()
        _set_paragraph_rtl(p)
        cap = d.get("caption", "") or ""
        itype = d.get("type", "") or ""
        product = d.get("product") or ""
        meta = " — ".join([x for x in [itype, (f"کالا: {product}" if product else ""), cap] if x])
        lead = p.add_run("🖼 تصویر: ")
        lead.bold = True
        linked = False
        if img_rel_dir:
            linked = _add_hyperlink(p, filename, f"{img_rel_dir}/{filename}")
        if not linked:
            p.add_run(filename)
        if meta:
            p.add_run(f"  ({meta})")


def render_markdown_into_doc(
    doc,
    md: str,
    page_images: Optional[Dict[int, List[tuple]]] = None,
    img_rel_dir: Optional[str] = None,
) -> None:
    """
    Markdown را داخل Document رندر می‌کند (RTL).
    اگر page_images داده شود، رفرنس تصاویرِ هر صفحه را در انتهای همان صفحه درج می‌کند.
    """
    lines = md.split("\n")
    i = 0
    n = len(lines)
    current_page: Optional[int] = None
    flushed_pages: set = set()

    while i < n:
        raw_line = lines[i]
        stripped = raw_line.strip()

        # نشانگر صفحه (خط مستقل)
        m_pg = re.fullmatch(r"<!--\s*page\s*:\s*(\d+)\s*-->", stripped)
        if m_pg:
            if current_page is not None and current_page not in flushed_pages:
                _flush_page_images(doc, page_images, current_page, img_rel_dir)
                flushed_pages.add(current_page)
            current_page = int(m_pg.group(1))
            i += 1
            continue

        # حذف نشانگرهای inline احتمالی از متن
        if "<!--" in stripped:
            stripped = _PAGE_MARK_RE.sub("", stripped).strip()
            if not stripped:
                i += 1
                continue

        if not stripped:
            i += 1
            continue

        # جدول
        if stripped.startswith("|") and i + 1 < n and _is_table_sep(lines[i + 1]):
            header = _parse_table_row(stripped)
            body: List[List[str]] = []
            j = i + 2
            while j < n and lines[j].strip().startswith("|"):
                body.append(_parse_table_row(lines[j]))
                j += 1
            _add_table(doc, header, body)
            i = j
            continue

        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            level = min(len(m.group(1)), 4)
            p = doc.add_heading(level=level)
            _set_paragraph_rtl(p)
            _add_md_runs(p, m.group(2).strip())
            i += 1
            continue

        m = re.match(r"^[-*+]\s+(.*)$", stripped)
        if m:
            p = doc.add_paragraph(style="List Bullet")
            _set_paragraph_rtl(p)
            _add_md_runs(p, m.group(1).strip())
            i += 1
            continue

        m = re.match(r"^\d+[.)]\s+(.*)$", stripped)
        if m:
            p = doc.add_paragraph(style="List Number")
            _set_paragraph_rtl(p)
            _add_md_runs(p, m.group(1).strip())
            i += 1
            continue

        p = doc.add_paragraph()
        _set_paragraph_rtl(p)
        _add_md_runs(p, stripped)
        i += 1

    # تصاویر آخرین صفحه
    if current_page is not None and current_page not in flushed_pages:
        _flush_page_images(doc, page_images, current_page, img_rel_dir)
        flushed_pages.add(current_page)

    # اگر هیچ نشانگر صفحه‌ای نبود ولی تصاویری مانده، همه را ته سند بگذار
    if page_images and not flushed_pages:
        for pg in sorted(page_images.keys()):
            _flush_page_images(doc, page_images, pg, img_rel_dir)


def add_images_appendix(doc, items: List[Dict[str, Any]], img_rel_dir: Optional[str]) -> None:
    """بخش «فهرست تصاویر» را به سند اضافه می‌کند (جدول: نام فایل، نوع، کالا، توضیح)."""
    if not items:
        return
    h = doc.add_heading(level=2)
    _set_paragraph_rtl(h)
    h.add_run("📷 فهرست تصاویر")

    header = ["نام فایل", "نوع", "کالا", "توضیح"]
    rows = []
    for it in items:
        rows.append([
            it.get("filename", ""),
            it.get("type", "") or "",
            it.get("product") or "",
            it.get("caption", "") or "",
        ])
    _add_table(doc, header, rows)


def _new_rtl_document(title: Optional[str] = None):
    doc = Document()

    # ── RTL سطح سند (body section) ──
    # بخش body → w:sectPr → bidi نیست، اما w:docDefaults همه‌ی استایل‌ها رو می‌گیره.
    # روش درست: تمام استایل‌هایی که متن دارن رو bidi کن + pPr پیش‌فرض سند رو bidi کن
    settings_element = doc.settings.element
    # w:rsid و w:writeProtection نگه‌می‌داریم؛ bidi روی docDefault می‌ذاریم
    try:
        docDefaults = settings_element.find(qn("w:docDefaults"))
        if docDefaults is None:
            docDefaults = OxmlElement("w:docDefaults")
            settings_element.insert(0, docDefaults)
        pPrDefault = docDefaults.find(qn("w:pPrDefault"))
        if pPrDefault is None:
            pPrDefault = OxmlElement("w:pPrDefault")
            docDefaults.append(pPrDefault)
        pPr = pPrDefault.find(qn("w:pPr"))
        if pPr is None:
            pPr = OxmlElement("w:pPr")
            pPrDefault.append(pPr)
        if pPr.find(qn("w:bidi")) is None:
            pPr.append(OxmlElement("w:bidi"))
        if pPr.find(qn("w:jc")) is None:
            jc = OxmlElement("w:jc"); jc.set(qn("w:val"), "right"); pPr.append(jc)
    except Exception:
        pass

    # ── فونت و RTL روی استایل Normal + تمام استایل‌های Paragraph ──
    def _style_rtl(style):
        try:
            pPr2 = style.element.get_or_add_pPr()
            if pPr2.find(qn("w:bidi")) is None:
                pPr2.append(OxmlElement("w:bidi"))
            pPr2.attrib.pop(qn("w:val"), None)
            style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        except Exception:
            pass

    from docx.enum.style import WD_STYLE_TYPE
    for s in doc.styles:
        if s.type == WD_STYLE_TYPE.PARAGRAPH:
            _style_rtl(s)

    normal = doc.styles["Normal"]
    normal.font.name = WORD_FONT
    normal.font.size = Pt(11)
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), WORD_FONT)
    rfonts.set(qn("w:hAnsi"), WORD_FONT)
    rfonts.set(qn("w:cs"), WORD_FONT)

    if title:
        h = doc.add_heading(level=0)
        _set_paragraph_rtl(h)
        h.add_run(title)
    return doc


def _build_page_images(
    occurrences: List[Dict[str, Any]],
    desc_map: Dict[str, Dict[str, Any]],
) -> Dict[int, List[tuple]]:
    page_images: Dict[int, List[tuple]] = {}
    seen_per_page: Dict[int, set] = {}
    for occ in occurrences:
        pg = occ["page"]
        fn = occ["filename"]
        seen = seen_per_page.setdefault(pg, set())
        if fn in seen:
            continue
        seen.add(fn)
        page_images.setdefault(pg, []).append((fn, desc_map.get(fn, {})))
    return page_images


def write_pdf_text_to_docx(
    md: str,
    out_path: Path,
    title: Optional[str] = None,
    page_images: Optional[Dict[int, List[tuple]]] = None,
    img_rel_dir: Optional[str] = None,
    appendix_items: Optional[List[Dict[str, Any]]] = None,
) -> None:
    doc = _new_rtl_document(title)
    render_markdown_into_doc(doc, md, page_images, img_rel_dir)
    if appendix_items:
        add_images_appendix(doc, appendix_items, img_rel_dir)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


# ---- Orchestrator ----

def generate_word_and_images(
    pdfs: List[Path],
    out_dir: Path,
    label: str,
    ts: str,
    key_manager: APIKeyManager,
    model: str,
    last_call_time: List[float],
    cancel_event: Optional[Event],
    log: Callable[[str], None],
    set_status: Callable[[str], None],
    subjects: Optional[List[str]] = None,
    desc_enabled: bool = True,
) -> None:
    if not pdfs:
        return
    subjects = subjects or []

    combined_doc = _new_rtl_document(f"متن کامل — {label}") if DOCX_AVAILABLE else None
    any_text = False

    log(f"\n  📝 استخراج متن و تصاویر ({len(pdfs)} فایل) ...")

    for idx, pdf in enumerate(pdfs, 1):
        if cancel_event is not None and cancel_event.is_set():
            log("⛔ لغو شد."); return
        set_status(f"متن/عکس: {pdf.name}")

        # هر PDF خروجی‌هایش داخل <stem>_out/ در کنار خودش ذخیره می‌شه
        pdf_out = make_pdf_out_dir(pdf)
        # فولدر عکس‌ها: <stem>_out/<stem>_out/  (همنام PDF با پسوند _out)
        img_dir = pdf_out / f"{pdf.stem}_out"
        img_dir.mkdir(parents=True, exist_ok=True)
        rel_dir = img_dir.name

        # --- تصاویر ---
        unique_records: List[Dict[str, Any]] = []
        occurrences: List[Dict[str, Any]] = []
        items: List[Dict[str, Any]] = []
        page_images: Dict[int, List[tuple]] = {}
        try:
            unique_records, occurrences = extract_images_from_pdf(pdf, img_dir)
            if unique_records:
                log(f"     🖼 {pdf.name}: {len(unique_records)} عکس یکتا")
                desc_map: Dict[str, Dict[str, Any]] = {}
                if desc_enabled:
                    set_status(f"توضیح عکس‌ها: {pdf.name}")
                    desc_map = describe_images_via_gemini(
                        unique_records, img_dir, key_manager, model,
                        last_call_time, cancel_event, subjects,
                    )
                else:
                    desc_map = {r["filename"]: {"type": "سایر", "product": None, "caption": ""}
                                for r in unique_records}
                for rec in unique_records:
                    embed_image_metadata(img_dir / rec["filename"], desc_map.get(rec["filename"], {}))
                items = write_images_index(img_dir, pdf.name, unique_records, occurrences, desc_map)
                page_images = _build_page_images(occurrences, desc_map)
            else:
                log(f"     🖼 {pdf.name}: عکسی یافت نشد")
        except Exception as e:
            log(f"     ⚠ خطا در عکس {pdf.name}: {str(e)[:60]}")

        if cancel_event is not None and cancel_event.is_set():
            log("⛔ لغو شد."); return

        # --- متن → Word (هم‌نام PDF با پسوند _out.docx) ---
        if not DOCX_AVAILABLE:
            continue
        try:
            md = extract_pdf_text_via_gemini(pdf, key_manager, model, last_call_time, cancel_event)
            if not md:
                log(f"     ✗ متن {pdf.name}: خروجی خالی")
                continue
            # نام فایل Word: <stem>_out.docx داخل <stem>_out/
            per_path = pdf_out / f"{pdf.stem}_out.docx"
            write_pdf_text_to_docx(
                md, per_path, title=pdf.stem,
                page_images=page_images, img_rel_dir=rel_dir, appendix_items=items,
            )
            if combined_doc is not None:
                sec = combined_doc.add_heading(level=1)
                _set_paragraph_rtl(sec)
                sec.add_run(pdf.stem)
                render_markdown_into_doc(combined_doc, md, page_images, rel_dir)
                add_images_appendix(combined_doc, items, rel_dir)
                combined_doc.add_page_break()
            any_text = True
            log(f"     📄 متن {pdf.name} ✓  →  {per_path.name}")
        except Exception as e:
            if cancel_event is not None and cancel_event.is_set():
                log("⛔ لغو شد."); return
            log(f"     ✗ خطا در متن {pdf.name}: {str(e)[:60]}")

    if DOCX_AVAILABLE and combined_doc is not None and any_text:
        try:
            combined_path = out_dir / f"{label}_out.docx"
            combined_doc.save(str(combined_path))
            log(f"  💾 Word ترکیبی: {combined_path.name}")
        except Exception as e:
            log(f"  ⚠ ذخیره Word ترکیبی ناموفق: {str(e)[:60]}")


# ======================== UI COLOR PALETTE ========================

C_ROOT   = "#011218"   # root background (darkest)
C_FRAME  = "#042e36"   # section frame background
C_TAB    = "#063f47"   # tab / section header base
C_TEAL   = "#0a6b77"   # accent teal / gradient end
C_BORDER = "#0d6b77"   # border highlight
C_GREEN  = "#0a8a5a"   # primary action (green)
C_MINT   = "#88ddcc"   # light label text
C_WHITE  = "#ffffff"   # primary text
C_LOG    = "#010d12"   # log area background
C_RED    = "#cc4444"   # danger / cancel
C_GRAY   = "#3a5a60"   # exit button

# legacy aliases used internally
UI_BG        = C_TAB
ACCENT_GREEN = C_GREEN
ACCENT_DARK  = C_TEAL


# ======================== UI ========================

def launch_ui(app_dir: Path) -> None:
    try:
        import tkinter as tk
        from tkinter import ttk, messagebox, filedialog, scrolledtext
    except ImportError as e:
        logging.warning(f"Tkinter unavailable: {e}")
        return

    config_path = app_dir / DEFAULT_CONFIG_NAME
    db_path = app_dir / "api_keys.db"
    cfg = load_config(config_path)

    root = tk.Tk()
    root.title("PDF Catalog Extractor")
    root.configure(bg=C_ROOT)
    root.minsize(660, 840)

    icon_path = get_resource_path("aa.ico")
    if icon_path.exists():
        try:
            root.iconbitmap(str(icon_path))
        except Exception:
            pass

    # ── TTK styles ──
    style = ttk.Style(root)
    try:
        style.theme_use("clam")
    except Exception:
        pass
    style.configure("TFrame", background=C_ROOT)
    style.configure("TLabel", background=C_ROOT, foreground=C_WHITE, font=("Arial", 10))
    style.configure("Sec.TFrame", background=C_FRAME)
    style.configure("DarkEntry.TEntry",
                    fieldbackground="#0a1820", foreground=C_WHITE,
                    insertcolor=C_WHITE, bordercolor=C_BORDER,
                    lightcolor=C_BORDER, darkcolor=C_FRAME, relief="flat")

    main_frame = tk.Frame(root, bg=C_ROOT)
    main_frame.pack(fill="both", expand=True)

    # ── Gradient helpers ──
    def _grad_v(canvas, top_hex: str, bot_hex: str, steps: int = 60):
        """Vertical gradient on canvas, top→bottom."""
        canvas.delete("g")
        w = canvas.winfo_width() or 700
        h = canvas.winfo_height() or 60
        r1, g1, b1 = [int(top_hex[i:i+2], 16) for i in (1, 3, 5)]
        r2, g2, b2 = [int(bot_hex[i:i+2], 16) for i in (1, 3, 5)]
        for i in range(steps):
            y0 = int(i * h / steps)
            y1 = int((i + 1) * h / steps)
            t = i / steps
            r = int(r1 + (r2 - r1) * t)
            g = int(g1 + (g2 - g1) * t)
            b = int(b1 + (b2 - b1) * t)
            canvas.create_rectangle(0, y0, w, y1 + 1,
                                    fill=f"#{r:02x}{g:02x}{b:02x}", outline="", tags="g")

    def _grad_h(canvas, left_hex: str, right_hex: str, steps: int = 50):
        """Horizontal gradient on canvas, left→right."""
        canvas.delete("g")
        w = canvas.winfo_width() or 700
        h = canvas.winfo_height() or 26
        r1, g1, b1 = [int(left_hex[i:i+2], 16) for i in (1, 3, 5)]
        r2, g2, b2 = [int(right_hex[i:i+2], 16) for i in (1, 3, 5)]
        for i in range(steps):
            x0 = int(i * w / steps)
            x1 = int((i + 1) * w / steps)
            t = i / steps
            r = int(r1 + (r2 - r1) * t)
            g = int(g1 + (g2 - g1) * t)
            b = int(b1 + (b2 - b1) * t)
            canvas.create_rectangle(x0, 0, x1 + 1, h,
                                    fill=f"#{r:02x}{g:02x}{b:02x}", outline="", tags="g")

    def make_section(parent, title: str):
        """Bordered card with a horizontal-gradient strip title."""
        outer = tk.Frame(parent, bg=C_BORDER, padx=1, pady=1)
        inner = tk.Frame(outer, bg=C_FRAME)
        inner.pack(fill="both", expand=True)
        hdr = tk.Canvas(inner, height=26, highlightthickness=0, bd=0)
        hdr.pack(fill="x")

        def _redraw(event=None, _c=hdr, _t=title):
            _grad_h(_c, C_TAB, C_TEAL, steps=40)
            _c.create_text(10, 13, text=_t, fill=C_WHITE,
                           font=("Arial", 9, "bold"), anchor="w", tags="g")

        hdr.bind("<Configure>", _redraw)
        body = tk.Frame(inner, bg=C_FRAME, padx=12, pady=6)
        body.pack(fill="both", expand=True)
        return outer, body

    # ── Banner header with vertical gradient ──
    banner = tk.Canvas(main_frame, height=76, highlightthickness=0, bd=0)
    banner.pack(fill="x")

    def _redraw_banner(event=None, _c=banner):
        _grad_v(_c, "#011218", "#0a6b77", steps=55)
        w = _c.winfo_width() or 700
        _c.create_line(0, 74, w, 74, fill=C_GREEN, width=2, tags="g")
        _c.create_text(w // 2, 28, text="PDF Catalog Extractor",
                       fill=C_WHITE, font=("Arial", 16, "bold"),
                       anchor="center", tags="g")
        _c.create_text(w // 2, 52, text="استخراج هوشمند اطلاعات از کاتالوگ‌های PDF",
                       fill=C_MINT, font=("Arial", 9),
                       anchor="center", tags="g")

    banner.bind("<Configure>", _redraw_banner)

    content = tk.Frame(main_frame, bg=C_ROOT, padx=14, pady=6)
    content.pack(fill="both", expand=True)

    # ── Top-level tab bar: اصلی / تنظیمات ──
    TOP_ACTIVE   = {"bg": C_GREEN,   "fg": C_WHITE, "relief": "flat",
                    "font": ("Arial", 11, "bold"), "bd": 0, "highlightthickness": 0}
    TOP_INACTIVE = {"bg": "#0a4a52", "fg": C_MINT,  "relief": "flat",
                    "font": ("Arial", 11), "bd": 0, "highlightthickness": 0}

    top_bar = tk.Frame(content, bg=C_ROOT)
    top_bar.pack(fill="x", pady=(0, 6))

    btn_top_main = tk.Button(top_bar, text="  🏠 اصلی  ", **TOP_ACTIVE,
                             cursor="hand2", padx=10, pady=6,
                             activebackground=C_GREEN, activeforeground=C_WHITE)
    btn_top_set  = tk.Button(top_bar, text="  ⚙ تنظیمات  ", **TOP_INACTIVE,
                             cursor="hand2", padx=10, pady=6,
                             activebackground=C_GREEN, activeforeground=C_WHITE)
    btn_top_main.pack(side="right", padx=(0, 2))
    btn_top_set.pack(side="right")

    top_pages = tk.Frame(content, bg=C_ROOT)
    top_pages.pack(fill="both", expand=True)

    settings_page = tk.Frame(top_pages, bg=C_ROOT)
    main_page = tk.Frame(top_pages, bg=C_ROOT)

    def switch_top(which: str):
        if which == "settings":
            main_page.pack_forget()
            settings_page.pack(fill="both", expand=True)
            btn_top_set.configure(**TOP_ACTIVE)
            btn_top_main.configure(**TOP_INACTIVE)
        else:
            settings_page.pack_forget()
            main_page.pack(fill="both", expand=True)
            btn_top_main.configure(**TOP_ACTIVE)
            btn_top_set.configure(**TOP_INACTIVE)

    btn_top_main.configure(command=lambda: switch_top("main"))
    btn_top_set.configure(command=lambda: switch_top("settings"))

    # ════════════════ تب تنظیمات ════════════════
    api_out, api_body = make_section(settings_page, "  Gemini API Keys")
    api_out.pack(fill="x", pady=(0, 8))

    tk.Label(api_body, text="هر خط یک کلید:", bg=C_FRAME, fg=C_MINT,
             font=("Arial", 9)).pack(anchor="w")
    keys_text = scrolledtext.ScrolledText(
        api_body, height=5, width=55,
        font=("Courier", 9), bg="#0a1820", fg="#00ff99",
        insertbackground=C_WHITE, relief="flat", bd=0,
    )
    keys_text.pack(fill="x", pady=(4, 0))
    saved_keys = cfg.get("gemini_api_keys", [])
    if isinstance(saved_keys, str):
        saved_keys = [saved_keys]
    if saved_keys:
        keys_text.insert("1.0", "\n".join(saved_keys))

    model_row = tk.Frame(api_body, bg=C_FRAME)
    model_row.pack(fill="x", pady=(8, 0))
    tk.Label(model_row, text="مدل Gemini:", bg=C_FRAME, fg=C_MINT,
             font=("Arial", 9)).pack(side="right")
    model_var = tk.StringVar(value=cfg.get("gemini_model", "gemini-2.0-flash"))
    ttk.Entry(model_row, textvariable=model_var, width=30,
              style="DarkEntry.TEntry").pack(side="right", padx=(8, 0))

    # گزینه‌ی توضیح‌گذاری تصاویر (vision)
    desc_var = tk.BooleanVar(value=cfg.get("image_desc_enabled", IMAGE_DESC_ENABLED_DEFAULT))
    tk.Checkbutton(
        settings_page,
        text="توضیح‌گذاری تصاویر با Gemini (یک API call اضافه به‌ازای هر PDF)",
        variable=desc_var, bg=C_ROOT, fg=C_MINT,
        selectcolor="#0a1820", activebackground=C_ROOT, activeforeground=C_WHITE,
        font=("Arial", 9), bd=0, highlightthickness=0,
    ).pack(anchor="e", pady=(8, 0))

    # ════════════════ تب اصلی ════════════════
    # ── موضوع (مولتی‌سلکت از روی subjects.json) ──
    subj_out, subj_body = make_section(main_page, "  موضوع (الزامی) — از لیست انتخاب کنید")
    subj_out.pack(fill="x", pady=(0, 8))

    subject_vars: Dict[str, Any] = {}
    selected_summary_var = tk.StringVar(value="موضوعی انتخاب نشده")

    subj_row = tk.Frame(subj_body, bg=C_FRAME)
    subj_row.pack(fill="x")

    subj_mb = tk.Menubutton(
        subj_row, textvariable=selected_summary_var,
        bg="#0a1820", fg=C_WHITE, font=("Arial", 10), relief="flat",
        anchor="e", padx=10, pady=6, cursor="hand2",
        activebackground=C_TEAL, activeforeground=C_WHITE, bd=0, highlightthickness=1,
        highlightbackground=C_BORDER,
    )
    subj_mb.pack(side="right", fill="x", expand=True)
    subj_menu = tk.Menu(subj_mb, tearoff=0, bg="#0a1820", fg=C_WHITE,
                        activebackground=C_TEAL, activeforeground=C_WHITE,
                        selectcolor=C_GREEN)
    subj_mb.config(menu=subj_menu)

    def _refresh_subject_summary():
        sel = [s for s, v in subject_vars.items() if v.get()]
        selected_summary_var.set("، ".join(sel) if sel else "موضوعی انتخاب نشده")

    def get_selected_subjects() -> List[str]:
        return [s for s, v in subject_vars.items() if v.get()]

    def _populate_subjects(items: List[str]):
        subj_menu.delete(0, "end")
        subject_vars.clear()
        last_sel = set(cfg.get("last_subjects", []))
        if items:
            for s in items:
                var = tk.BooleanVar(value=(s in last_sel))
                subject_vars[s] = var
                subj_menu.add_checkbutton(label=s, variable=var,
                                          command=_refresh_subject_summary)
        else:
            subj_menu.add_command(label="— لیستی بارگیری نشد (اتصال/پیکربندی) —",
                                  state="disabled")
        _refresh_subject_summary()

    def _reload_subjects():
        items = load_subjects(cfg)
        save_config(config_path, cfg)
        _populate_subjects(items)
        set_status(f"لیست موضوع‌ها: {len(items)} مورد" if items else "لیست موضوع خالی")

    _mk_subj_btn = lambda txt, cmd: tk.Button(
        subj_row, text=txt, command=cmd, bg=C_TEAL, fg=C_WHITE, font=("Arial", 9),
        relief="flat", cursor="hand2", padx=8, pady=4,
        activebackground=C_GREEN, activeforeground=C_WHITE, bd=0, highlightthickness=0)
    _mk_subj_btn("🔄 بارگیری مجدد", _reload_subjects).pack(side="left", padx=(6, 0))

    # ── Inputs ──
    in_out, in_body = make_section(main_page, "  ورودی — فولدر یا PDF (هر خط یک مسیر)")
    in_out.pack(fill="x", pady=(0, 8))

    inputs_text = scrolledtext.ScrolledText(
        in_body, height=4, width=55,
        font=("Arial", 9), bg="#0a1820", fg=C_WHITE, wrap="none",
        insertbackground=C_WHITE, relief="flat", bd=0,
    )
    inputs_text.pack(fill="x", pady=(0, 6))
    if cfg.get("last_inputs"):
        inputs_text.insert("1.0", "\n".join(cfg["last_inputs"]))

    btn_in_row = tk.Frame(in_body, bg=C_FRAME)
    btn_in_row.pack(fill="x")

    def _mk_btn(parent, text, cmd, bg=None):
        return tk.Button(parent, text=text, command=cmd,
                         bg=bg or C_TEAL, fg=C_WHITE, font=("Arial", 9),
                         relief="flat", cursor="hand2", padx=8, pady=4,
                         activebackground=C_GREEN, activeforeground=C_WHITE,
                         bd=0, highlightthickness=0)

    def add_folder():
        path = filedialog.askdirectory(title="انتخاب فولدر")
        if path:
            cur = inputs_text.get("1.0", "end").strip()
            inputs_text.insert("end", ("" if cur == "" else "\n") + path)

    def add_pdfs():
        paths = filedialog.askopenfilenames(
            title="انتخاب PDF", filetypes=[("PDF files", "*.pdf")]
        )
        for p in paths:
            cur = inputs_text.get("1.0", "end").strip()
            inputs_text.insert("end", ("" if cur == "" else "\n") + p)

    _mk_btn(btn_in_row, "➕ افزودن فولدر", add_folder).pack(side="left", padx=(0, 6))
    _mk_btn(btn_in_row, "➕ افزودن PDF", add_pdfs).pack(side="left")

    # ── Sub-tab bar: خودکار / با مدل ──
    tab_bar = tk.Frame(main_page, bg=C_ROOT)
    tab_bar.pack(fill="x", pady=(2, 0))

    TAB_ACTIVE   = {"bg": C_GREEN,   "fg": C_WHITE, "relief": "flat",
                    "font": ("Arial", 10, "bold"), "bd": 0, "highlightthickness": 0}
    TAB_INACTIVE = {"bg": "#0a4a52", "fg": C_MINT,  "relief": "flat",
                    "font": ("Arial", 10), "bd": 0, "highlightthickness": 0}

    mode_var = tk.StringVar(value=cfg.get("last_mode", "auto"))

    tab_model_btn = tk.Button(tab_bar, text="  🔍 با مدل  ", **TAB_INACTIVE,
                              cursor="hand2", padx=6, pady=5,
                              activebackground=C_GREEN, activeforeground=C_WHITE)
    tab_auto_btn  = tk.Button(tab_bar, text="  📁 خودکار  ", **TAB_INACTIVE,
                              cursor="hand2", padx=6, pady=5,
                              activebackground=C_GREEN, activeforeground=C_WHITE)
    tab_model_btn.pack(side="right", padx=(0, 2))
    tab_auto_btn.pack(side="right")

    tab_content = tk.Frame(main_page, bg=C_FRAME)
    tab_content.pack(fill="both", expand=True, pady=(2, 0))

    # ── Panel: جستجو با مدل ──
    panel_model = tk.Frame(tab_content, bg=C_FRAME)

    def _dark_text(parent, height):
        return scrolledtext.ScrolledText(
            parent, height=height, width=55,
            font=("Arial", 10), bg="#0a1820", fg=C_WHITE, wrap="word",
            insertbackground=C_WHITE, relief="flat", bd=0,
        )

    def _sec_lbl(parent, text):
        tk.Label(parent, text=text, bg=C_FRAME, fg=C_MINT,
                 font=("Arial", 9)).pack(anchor="w", padx=8, pady=(6, 0))

    _sec_lbl(panel_model, "مدل‌ها — هر خط یک مدل — اختیاری (خالی = همه):")
    models_text = _dark_text(panel_model, 5)
    models_text.pack(fill="both", expand=True, padx=8, pady=(2, 4))
    if cfg.get("last_models"):
        models_text.insert("1.0", "\n".join(cfg["last_models"]))

    _sec_lbl(panel_model, "ویژگی‌ها — هر خط یک ویژگی — اجباری:")
    feat1_text = _dark_text(panel_model, 5)
    feat1_text.pack(fill="both", expand=True, padx=8, pady=(2, 4))
    if cfg.get("last_features"):
        feat1_text.insert("1.0", "\n".join(cfg["last_features"]))

    # ── Panel: خودکار / فولدر ──
    panel_auto = tk.Frame(tab_content, bg=C_FRAME)

    cat_row = tk.Frame(panel_auto, bg=C_FRAME)
    cat_row.pack(fill="x", padx=8, pady=(8, 4))
    tk.Label(cat_row, text="نوع کالا (اختیاری):", bg=C_FRAME, fg=C_MINT,
             font=("Arial", 9)).pack(side="left")
    category_var = tk.StringVar(value=cfg.get("last_category", ""))
    ttk.Entry(cat_row, textvariable=category_var, width=36,
              style="DarkEntry.TEntry").pack(side="left", padx=(8, 0))

    manual_feat_var = tk.BooleanVar(value=cfg.get("manual_features_enabled", False))
    tk.Checkbutton(
        panel_auto,
        text="ویژگی‌ها را خودم دستی وارد می‌کنم (در غیر این صورت از MERG.pdf کشف می‌شود)",
        variable=manual_feat_var, bg=C_FRAME, fg=C_MINT,
        selectcolor="#0a1820", activebackground=C_FRAME, activeforeground=C_WHITE,
        font=("Arial", 9), bd=0, highlightthickness=0,
    ).pack(anchor="w", padx=8, pady=(0, 4))

    feat2_frame = tk.Frame(panel_auto, bg=C_FRAME)
    _sec_lbl(feat2_frame, "ویژگی‌های دستی — هر خط یک ویژگی:")
    feat2_text = _dark_text(feat2_frame, 8)
    feat2_text.pack(fill="both", expand=True, padx=8, pady=(2, 4))
    if cfg.get("last_manual_features"):
        feat2_text.insert("1.0", "\n".join(cfg["last_manual_features"]))

    def on_manual_feat_toggle(*_):
        if manual_feat_var.get():
            feat2_frame.pack(fill="both", expand=True)
        else:
            feat2_frame.pack_forget()

    manual_feat_var.trace_add("write", on_manual_feat_toggle)
    on_manual_feat_toggle()

    def switch_tab(tab: str):
        mode_var.set(tab)
        if tab == "model":
            panel_auto.pack_forget()
            panel_model.pack(fill="both", expand=True)
            tab_model_btn.configure(**TAB_ACTIVE)
            tab_auto_btn.configure(**TAB_INACTIVE)
        else:
            panel_model.pack_forget()
            panel_auto.pack(fill="both", expand=True)
            tab_model_btn.configure(**TAB_INACTIVE)
            tab_auto_btn.configure(**TAB_ACTIVE)

    tab_model_btn.configure(command=lambda: switch_tab("model"))
    tab_auto_btn.configure(command=lambda: switch_tab("auto"))

    if cfg.get("last_mode", "auto") == "model":
        switch_tab("model")
    else:
        switch_tab("auto")

    # ── Status bar ──
    status_var = tk.StringVar(value="آماده")
    tk.Label(content, textvariable=status_var, bg=C_ROOT, fg=C_MINT,
             font=("Arial", 9), anchor="w").pack(fill="x", pady=(4, 0))

    # ── Action buttons ──
    btn_frame = tk.Frame(content, bg=C_ROOT)
    btn_frame.pack(fill="x", pady=(6, 4))

    # ── Log section ──
    log_out = tk.Frame(content, bg=C_BORDER, padx=1, pady=1)
    log_out.pack(fill="both", expand=True, pady=(4, 0))
    log_wrap = tk.Frame(log_out, bg=C_LOG)
    log_wrap.pack(fill="both", expand=True)

    log_hdr = tk.Canvas(log_wrap, height=24, highlightthickness=0, bd=0)
    log_hdr.pack(fill="x")

    def _redraw_log_hdr(event=None, _c=log_hdr):
        _grad_h(_c, C_LOG, C_TAB, steps=30)
        _c.create_text(10, 12, text="📋 لاگ", fill=C_MINT,
                       font=("Arial", 9, "bold"), anchor="w", tags="g")

    log_hdr.bind("<Configure>", _redraw_log_hdr)

    log_text = scrolledtext.ScrolledText(
        log_wrap, height=10, bg=C_LOG, fg="#aaffcc",
        font=("Consolas", 9), state="disabled", wrap="word",
        relief="flat", bd=0, insertbackground=C_WHITE,
    )
    log_text.pack(fill="both", expand=True, padx=4, pady=(0, 4))

    def log(msg: str):
        def _u():
            log_text.config(state="normal")
            log_text.insert("end", msg + "\n")
            log_text.see("end")
            log_text.config(state="disabled")
        root.after(0, _u)

    def set_status(msg: str):
        root.after(0, lambda: status_var.set(msg))

    # مقداردهی اولیه‌ی لیست موضوع‌ها (از remote یا کش) و انتخاب تب اصلی
    try:
        _initial_subjects = load_subjects(cfg)
        save_config(config_path, cfg)
        _populate_subjects(_initial_subjects)
    except Exception as _e:
        logging.debug(f"subjects init failed: {_e}")
        _populate_subjects([])
    switch_top("main")

    def apply_features_to_manual(features: List[str]):
        def _u():
            feat2_text.delete("1.0", "end")
            if features:
                feat2_text.insert("1.0", "\n".join(features))
            manual_feat_var.set(True)
        root.after(0, _u)

    cancel_event = Event()
    runtime_opts: Dict[str, Any] = {"desc_enabled": True}

    def _desc_enabled() -> bool:
        return bool(runtime_opts.get("desc_enabled", True))

    def _post_run_notify(
        cfg: Dict[str, Any],
        subjects: List[str],
        ts: str,
        input_pdfs: List[Path],
        out_dirs: List[tuple],
        total_rows: int,
        log: Callable[[str], None],
    ):
        """آپلود به Drive و ارسال پیام تلگرام — در thread پس‌زمینه اجرا می‌شود."""
        if not GDRIVE_AVAILABLE:
            _tgram_notify(cfg, "کاربر", subjects, ts, success=True, drive_links=None)
            return

        drive_root = (DRIVE_ROOT_FOLDER_NAME.strip() or "CatalogExtractor")
        token_path = app_dir / "drive_token.json"
        drive_links: Optional[Dict[str, str]] = None
        err_msg = ""

        try:
            service = drive_build_service(token_path)
            # ساخت فولدر اجرا در Drive و آپلود همه‌ی خروجی‌ها
            subj_label = "، ".join(subjects) if subjects else "عمومی"
            root_id = _drive_get_or_create_folder(service, drive_root)
            subj_id = _drive_get_or_create_folder(service, subj_label, root_id)
            ts_id   = _drive_get_or_create_folder(service, ts, subj_id)

            links_out: Dict[str, str] = {}
            for (in_dir, out_dir_path) in out_dirs:
                # فولدر ورودی
                in_drive_id  = _drive_get_or_create_folder(service, in_dir.name + "_ورودی", ts_id)
                # آپلود PDFهای این ورودی
                for pdf in input_pdfs:
                    if pdf.parent == in_dir:
                        try:
                            _drive_upload_file(service, pdf, in_drive_id)
                            log(f"     ☁ آپلود ورودی: {pdf.name}")
                        except Exception as e:
                            log(f"     ⚠ آپلود ناموفق ({pdf.name}): {str(e)[:50]}")
                # فولدر خروجی
                out_drive_id = _drive_get_or_create_folder(service, out_dir_path.name, ts_id)
                def _up_recursive(local: Path, pid: str):
                    for item in sorted(local.iterdir()):
                        if item.is_file():
                            try:
                                _drive_upload_file(service, item, pid)
                                log(f"     ☁ آپلود خروجی: {item.name}")
                            except Exception as e:
                                log(f"     ⚠ آپلود ناموفق ({item.name}): {str(e)[:50]}")
                        elif item.is_dir():
                            sub = _drive_get_or_create_folder(service, item.name, pid)
                            _up_recursive(item, sub)
                _up_recursive(out_dir_path, out_drive_id)
                links_out[out_dir_path.name] = f"https://drive.google.com/drive/folders/{out_drive_id}"

            drive_links = {
                "run_folder_link": f"https://drive.google.com/drive/folders/{ts_id}",
                "output_folder_link": f"https://drive.google.com/drive/folders/{ts_id}",
            }
            if links_out:
                first_key = next(iter(links_out))
                drive_links["output_folder_link"] = links_out[first_key]

            log(f"\n  ☁ آپلود Drive کامل شد.")
            _tgram_notify(cfg, "کاربر", subjects, ts, success=True, drive_links=drive_links)

        except Exception as e:
            err_msg = str(e)[:200]
            log(f"\n  ⚠ آپلود Drive ناموفق: {err_msg}")
            _tgram_notify(cfg, "کاربر", subjects, ts, success=False, error_msg=err_msg)

    # ── Feature confirmation dialog ──
    def ask_features_confirm(folder_name: str, features: List[str]) -> Optional[tuple]:
        result_holder: Dict[str, Any] = {"value": None, "done": Event()}

        def _build():
            import tkinter as tk
            from tkinter import scrolledtext as st
            dlg = tk.Toplevel(root)
            dlg.title(f"تأیید ویژگی‌ها — {folder_name}")
            dlg.configure(bg=C_ROOT)
            dlg.transient(root)
            dlg.grab_set()
            dlg.minsize(440, 460)

            icon_p = get_resource_path("aa.ico")
            if icon_p.exists():
                try:
                    dlg.iconbitmap(str(icon_p))
                except Exception:
                    pass

            dlg_hdr = tk.Canvas(dlg, height=62, highlightthickness=0, bd=0)
            dlg_hdr.pack(fill="x")

            def _draw_dlg_hdr(event=None, _c=dlg_hdr):
                _grad_v(_c, "#011218", "#0a6b77", steps=40)
                w = _c.winfo_width() or 440
                _c.create_text(w // 2, 22, text="ویژگی‌های کشف‌شده",
                               fill=C_WHITE, font=("Arial", 12, "bold"),
                               anchor="center", tags="g")
                _c.create_text(w // 2, 44, text=folder_name,
                               fill=C_MINT, font=("Arial", 9),
                               anchor="center", tags="g")
                _c.create_line(0, 60, w, 60, fill=C_GREEN, width=1, tags="g")

            dlg_hdr.bind("<Configure>", _draw_dlg_hdr)

            tk.Label(dlg, text="می‌توانید ویرایش کنید، حذف کنید یا خط جدید اضافه کنید (هر خط یک ویژگی):",
                     bg=C_ROOT, fg=C_MINT, font=("Arial", 8)).pack(padx=12, anchor="w", pady=(8, 2))

            txt = st.ScrolledText(dlg, height=12, width=48, font=("Arial", 10),
                                  bg="#0a1820", fg=C_WHITE, wrap="word",
                                  relief="flat", bd=0, insertbackground=C_WHITE)
            txt.pack(fill="both", expand=True, padx=12, pady=(0, 8))
            if features:
                txt.insert("1.0", "\n".join(features))

            bar = tk.Frame(dlg, bg=C_ROOT)
            bar.pack(fill="x", padx=12, pady=(0, 12))

            def confirm():
                feats = [f.strip() for f in txt.get("1.0", "end").splitlines() if f.strip()]
                result_holder["value"] = ("confirm", feats)
                dlg.destroy()
                result_holder["done"].set()

            def to_manual():
                feats = [f.strip() for f in txt.get("1.0", "end").splitlines() if f.strip()]
                result_holder["value"] = ("to_manual", feats)
                dlg.destroy()
                result_holder["done"].set()

            def cancel_dlg():
                result_holder["value"] = None
                dlg.destroy()
                result_holder["done"].set()

            tk.Button(bar, text="✅ تأیید و ادامه", command=confirm,
                      bg=C_GREEN, fg=C_WHITE, font=("Arial", 10, "bold"),
                      relief="flat", cursor="hand2", padx=14, pady=6,
                      activebackground="#0daa6a", bd=0).pack(side="left")
            tk.Button(bar, text="➡️ انتقال به ورود دستی", command=to_manual,
                      bg=C_TEAL, fg=C_WHITE, font=("Arial", 9),
                      relief="flat", cursor="hand2", padx=10, pady=6,
                      activebackground=C_GREEN, bd=0).pack(side="left", padx=(8, 0))
            tk.Button(bar, text="لغو کامل", command=cancel_dlg,
                      bg=C_RED, fg=C_WHITE, font=("Arial", 9),
                      relief="flat", cursor="hand2", padx=10, pady=6,
                      activebackground="#ee5555", bd=0).pack(side="left", padx=(8, 0))

            dlg.protocol("WM_DELETE_WINDOW", cancel_dlg)

        root.after(0, _build)
        result_holder["done"].wait()
        return result_holder["value"]

    # ── worker: تب «خودکار / فولدر» ──
    def run_auto_mode(api_keys, model_name, inputs, category, manual_features, subjects=None):
        last_call_time = [0.0]
        subjects = subjects or []
        if not category.strip() and subjects:
            category = "، ".join(subjects)
        try:
            key_manager = APIKeyManager(api_keys, db_path=db_path)
            if key_manager.active_count == 0:
                log("✗ هیچ API Key فعالی نیست.")
                set_status("✗ کلید فعال موجود نیست")
                return

            folders: List[Path] = []
            single_pdfs: List[Path] = []
            for raw in inputs:
                p = Path(raw)
                if p.is_dir():
                    folders.append(p)
                elif p.is_file() and p.suffix.lower() == ".pdf":
                    if not is_merg(p):
                        single_pdfs.append(p)
                else:
                    log(f"⚠ نادیده گرفته شد (نامعتبر): {raw}")

            ts = datetime.now().strftime("%Y%m%d_%H%M%S")
            manual_features = list(manual_features)
            total_rows = 0
            last_out_dir: Optional[Path] = None
            all_input_pdfs: List[Path] = []
            all_out_dirs: List[tuple] = []   # (input_dir, out_dir)

            for folder in folders:
                if cancel_event.is_set():
                    log("⛔ لغو شد."); set_status("لغو شد"); return

                log(f"\n📁 فولدر: {folder.name}")
                set_status(f"فولدر: {folder.name}")
                pdfs = collect_folder_pdfs(folder)
                merg = find_merg(folder)

                if not pdfs:
                    log("  ⚠ هیچ PDF (غیر MERG) در این فولدر نیست — رد شد.")
                    continue

                if manual_features:
                    features = list(manual_features)
                    log(f"  ✓ ویژگی‌های دستی استفاده می‌شوند: {', '.join(features)}")
                else:
                    features = []
                    if merg:
                        log("  🔎 کشف ویژگی‌ها از MERG.pdf ...")
                        set_status("کشف ویژگی‌ها...")
                        try:
                            features = discover_features(merg, category, key_manager, model_name, last_call_time, cancel_event, log)
                            log(f"  ✓ {len(features)} ویژگی کشف شد.")
                        except Exception as e:
                            if cancel_event.is_set():
                                log("⛔ لغو شد."); set_status("لغو شد"); return
                            log(f"  ✗ خطا در کشف ویژگی: {str(e)[:80]}")
                            features = []
                    else:
                        log("  ℹ MERG.pdf نیست — ویژگی‌ها را دستی وارد کنید.")

                    if cancel_event.is_set():
                        log("⛔ لغو شد."); set_status("لغو شد"); return

                    set_status("منتظر تأیید ویژگی‌ها...")
                    decision = ask_features_confirm(folder.name, features)
                    if decision is None:
                        log("  ⛔ لغو کامل توسط کاربر."); set_status("لغو شد"); return

                    action, feats = decision
                    features = feats
                    if action == "to_manual":
                        manual_features = list(feats)
                        apply_features_to_manual(manual_features)
                        log(f"  ➡️ ویژگی‌ها به ورود دستی منتقل شد: {', '.join(features)}")
                    else:
                        log(f"  ✓ ویژگی‌های نهایی: {', '.join(features)}")

                if not features:
                    log("  ⚠ هیچ ویژگی‌ای مشخص نشد — این فولدر رد شد.")
                    continue

                folder_results: List[Dict[str, Any]] = []
                for idx, pdf in enumerate(pdfs, 1):
                    if cancel_event.is_set():
                        log("⛔ لغو شد."); set_status("لغو شد"); return
                    log(f"  📄 ({idx}/{len(pdfs)}) {pdf.name} ...")
                    set_status(f"{folder.name}: {pdf.name}")
                    try:
                        rows = extract_pdf(pdf, features, category, key_manager, model_name, last_call_time, cancel_event)
                        folder_results.extend(rows)
                        log(f"     ✓ {len(rows)} ردیف")
                    except Exception as e:
                        if cancel_event.is_set():
                            log("⛔ لغو شد."); set_status("لغو شد"); return
                        log(f"     ✗ خطا: {str(e)[:80]}")

                folder_out = make_folder_out_dir(folder, ts)
                if folder_results:
                    fout = folder_out / f"{folder.name}_out.xlsx"
                    create_excel_output(folder_results, features, fout, include_source=True)
                    log(f"  💾 ذخیره: {fout.name}  ({len(folder_results)} ردیف)")
                    total_rows += len(folder_results)
                    last_out_dir = folder_out
                else:
                    log("  ⚠ نتیجه‌ای برای این فولدر نبود.")

                # متن (Word جدا + ترکیبی) و تصاویر برای کل PDFهای این فولدر
                generate_word_and_images(
                    pdfs, folder_out, folder.name, ts,
                    key_manager, model_name, last_call_time, cancel_event, log, set_status, subjects=subjects, desc_enabled=_desc_enabled(),
                )
                last_out_dir = folder_out
                all_input_pdfs.extend(pdfs)
                all_out_dirs.append((folder, folder_out))

            if single_pdfs:
                if cancel_event.is_set():
                    log("⛔ لغو شد."); set_status("لغو شد"); return
                log(f"\n📄 PDFهای تکی: {len(single_pdfs)} فایل")

                if manual_features:
                    features = list(manual_features)
                    log(f"  ✓ ویژگی‌های دستی استفاده می‌شوند: {', '.join(features)}")
                    action, feats = "confirm", features
                else:
                    log("  🔎 کشف ویژگی‌ها از اولین PDF ...")
                    set_status("کشف ویژگی‌ها (PDF تکی)...")
                    try:
                        feats0 = discover_features(single_pdfs[0], category, key_manager, model_name, last_call_time, cancel_event, log)
                        log(f"  ✓ {len(feats0)} ویژگی کشف شد.")
                    except Exception as e:
                        if cancel_event.is_set():
                            log("⛔ لغو شد."); set_status("لغو شد"); return
                        log(f"  ✗ خطا در کشف ویژگی: {str(e)[:80]}")
                        feats0 = []

                    if cancel_event.is_set():
                        log("⛔ لغو شد."); set_status("لغو شد"); return

                    decision = ask_features_confirm("PDFهای تکی", feats0)
                    if decision is None:
                        log("  ⛔ لغو کامل."); set_status("لغو شد"); return
                    action, feats = decision

                features = feats
                if action == "to_manual":
                    manual_features = list(feats)
                    apply_features_to_manual(manual_features)
                    log(f"  ➡️ ویژگی‌ها به ورود دستی منتقل شد: {', '.join(features)}")

                if features:
                    log(f"  ✓ ویژگی‌های نهایی: {', '.join(features)}")
                    single_results: List[Dict[str, Any]] = []
                    for idx, pdf in enumerate(single_pdfs, 1):
                        if cancel_event.is_set():
                            log("⛔ لغو شد."); set_status("لغو شد"); return
                        log(f"  📄 ({idx}/{len(single_pdfs)}) {pdf.name} ...")
                        set_status(f"PDF تکی: {pdf.name}")
                        try:
                            rows = extract_pdf(pdf, features, category, key_manager, model_name, last_call_time, cancel_event)
                            single_results.extend(rows)
                            log(f"     ✓ {len(rows)} ردیف")
                        except Exception as e:
                            if cancel_event.is_set():
                                log("⛔ لغو شد."); set_status("لغو شد"); return
                            log(f"     ✗ خطا: {str(e)[:80]}")

                    # برای هر PDF تکی یک فولدر _out جداگانه می‌سازیم
                    for sp in single_pdfs:
                        sp_out = make_pdf_out_dir(sp)
                        sp_results = [r for r in single_results if r.get("_source") == sp.name]
                        if sp_results:
                            fout = sp_out / f"{sp.stem}_out.xlsx"
                            create_excel_output(sp_results, features, fout, include_source=False)
                            log(f"  💾 ذخیره: {fout.name}")
                        all_input_pdfs.append(sp)
                        all_out_dirs.append((sp.parent, sp_out))
                    if single_results:
                        total_rows += len(single_results)
                        last_out_dir = make_pdf_out_dir(single_pdfs[-1])

                    # متن و تصاویر (هر PDF → فولدر _out خودش)
                    generate_word_and_images(
                        single_pdfs, single_pdfs[0].parent, "PDFهای_تکی", ts,
                        key_manager, model_name, last_call_time, cancel_event, log, set_status, subjects=subjects, desc_enabled=_desc_enabled(),
                    )
                else:
                    log("  ⚠ هیچ ویژگی‌ای مشخص نشد — PDFهای تکی رد شد.")

            _post_run_notify(cfg, subjects, ts, all_input_pdfs, all_out_dirs, total_rows, log)

            if total_rows:
                log(f"\n✅ تمام شد — مجموعاً {total_rows} ردیف.")
                set_status(f"✅ تمام — {total_rows} ردیف")
                if last_out_dir:
                    try:
                        import subprocess
                        subprocess.Popen(["explorer", str(last_out_dir)])
                    except Exception:
                        pass
            else:
                log("\n⚠ هیچ نتیجه‌ای تولید نشد.")
                set_status("بدون نتیجه")

        except Exception as e:
            log(f"✗ خطای کلی: {str(e)[:100]}")
            logging.error(f"Fatal in run_auto_mode: {e}", exc_info=True)
            set_status(f"✗ خطا: {str(e)[:50]}")
        finally:
            root.after(0, lambda: start_btn.config(state="normal"))
            root.after(0, lambda: cancel_btn.pack_forget())

    # ── worker: تب «جستجو با مدل» ──
    def run_model_mode(api_keys, model_name, inputs, models, features, subjects=None):
        last_call_time = [0.0]
        subjects = subjects or []
        mode_category = "، ".join(subjects) if subjects else ""
        try:
            key_manager = APIKeyManager(api_keys, db_path=db_path)
            if key_manager.active_count == 0:
                log("✗ هیچ API Key فعالی نیست.")
                set_status("✗ کلید فعال موجود نیست")
                return

            folders: List[Path] = []
            single_pdfs: List[Path] = []
            for raw in inputs:
                p = Path(raw)
                if p.is_dir():
                    folders.append(p)
                elif p.is_file() and p.suffix.lower() == ".pdf":
                    if not is_merg(p):
                        single_pdfs.append(p)
                else:
                    log(f"⚠ نادیده گرفته شد (نامعتبر): {raw}")

            ts = datetime.now().strftime("%Y%m%d_%H%M%S")
            total_rows = 0
            last_out_dir: Optional[Path] = None
            all_input_pdfs: List[Path] = []
            all_out_dirs: List[tuple] = []

            for folder in folders:
                if cancel_event.is_set():
                    log("⛔ لغو شد."); set_status("لغو شد"); return
                log(f"\n📁 فولدر: {folder.name}")
                pdfs = collect_folder_pdfs(folder)
                if not pdfs:
                    log("  ⚠ هیچ PDF (غیر MERG) در این فولدر نیست — رد شد.")
                    continue

                folder_results: List[Dict[str, Any]] = []
                for idx, pdf in enumerate(pdfs, 1):
                    if cancel_event.is_set():
                        log("⛔ لغو شد."); set_status("لغو شد"); return
                    log(f"  📄 ({idx}/{len(pdfs)}) {pdf.name} ...")
                    set_status(f"{folder.name}: {pdf.name}")
                    try:
                        rows = extract_pdf_model_mode(pdf, models, features, key_manager, model_name, last_call_time, cancel_event, category=mode_category)
                        folder_results.extend(rows)
                        log(f"     ✓ {len(rows)} ردیف")
                    except Exception as e:
                        if cancel_event.is_set():
                            log("⛔ لغو شد."); set_status("لغو شد"); return
                        log(f"     ✗ خطا: {str(e)[:80]}")

                folder_out = make_folder_out_dir(folder, ts)
                if folder_results:
                    fout = folder_out / f"{folder.name}_out.xlsx"
                    create_excel_output(folder_results, features, fout, include_source=True)
                    log(f"  💾 ذخیره: {fout.name}  ({len(folder_results)} ردیف)")
                    total_rows += len(folder_results)
                    last_out_dir = folder_out
                else:
                    log("  ⚠ نتیجه‌ای برای این فولدر نبود.")

                generate_word_and_images(
                    pdfs, folder_out, folder.name, ts,
                    key_manager, model_name, last_call_time, cancel_event, log, set_status, subjects=subjects, desc_enabled=_desc_enabled(),
                )
                last_out_dir = folder_out
                all_input_pdfs.extend(pdfs)
                all_out_dirs.append((folder, folder_out))

            if single_pdfs:
                if cancel_event.is_set():
                    log("⛔ لغو شد."); set_status("لغو شد"); return
                log(f"\n📄 PDFهای تکی: {len(single_pdfs)} فایل")
                single_results: List[Dict[str, Any]] = []
                for idx, pdf in enumerate(single_pdfs, 1):
                    if cancel_event.is_set():
                        log("⛔ لغو شد."); set_status("لغو شد"); return
                    log(f"  📄 ({idx}/{len(single_pdfs)}) {pdf.name} ...")
                    set_status(f"PDF تکی: {pdf.name}")
                    try:
                        rows = extract_pdf_model_mode(pdf, models, features, key_manager, model_name, last_call_time, cancel_event, category=mode_category)
                        single_results.extend(rows)
                        log(f"     ✓ {len(rows)} ردیف")
                    except Exception as e:
                        if cancel_event.is_set():
                            log("⛔ لغو شد."); set_status("لغو شد"); return
                        log(f"     ✗ خطا: {str(e)[:80]}")

                for sp in single_pdfs:
                    sp_out = make_pdf_out_dir(sp)
                    sp_results = [r for r in single_results if r.get("_source") == sp.name]
                    if sp_results:
                        fout = sp_out / f"{sp.stem}_out.xlsx"
                        create_excel_output(sp_results, features, fout, include_source=False)
                        log(f"  💾 ذخیره: {fout.name}")
                    all_input_pdfs.append(sp)
                    all_out_dirs.append((sp.parent, sp_out))
                if single_results:
                    total_rows += len(single_results)
                    last_out_dir = make_pdf_out_dir(single_pdfs[-1])

                generate_word_and_images(
                    single_pdfs, single_pdfs[0].parent, "PDFهای_تکی", ts,
                    key_manager, model_name, last_call_time, cancel_event, log, set_status, subjects=subjects, desc_enabled=_desc_enabled(),
                )

            _post_run_notify(cfg, subjects, ts, all_input_pdfs, all_out_dirs, total_rows, log)

            if total_rows:
                log(f"\n✅ تمام شد — مجموعاً {total_rows} ردیف.")
                set_status(f"✅ تمام — {total_rows} ردیف")
                if last_out_dir:
                    try:
                        import subprocess
                        subprocess.Popen(["explorer", str(last_out_dir)])
                    except Exception:
                        pass
            else:
                log("\n⚠ هیچ نتیجه‌ای تولید نشد.")
                set_status("بدون نتیجه")

        except Exception as e:
            log(f"✗ خطای کلی: {str(e)[:100]}")
            logging.error(f"Fatal in run_model_mode: {e}", exc_info=True)
            set_status(f"✗ خطا: {str(e)[:50]}")
        finally:
            root.after(0, lambda: start_btn.config(state="normal"))
            root.after(0, lambda: cancel_btn.pack_forget())

    def on_start():
        raw_keys = keys_text.get("1.0", "end").strip()
        api_keys = [k.strip() for k in raw_keys.splitlines() if k.strip()]
        model_name = model_var.get().strip() or "gemini-2.0-flash"
        inputs = [x.strip() for x in inputs_text.get("1.0", "end").splitlines() if x.strip()]
        current_mode = mode_var.get()

        if not api_keys:
            messagebox.showerror("خطا", "حداقل یک API Key وارد کنید."); return
        if not inputs:
            messagebox.showerror("خطا", "حداقل یک فولدر یا PDF وارد کنید."); return

        subjects = get_selected_subjects()
        if not subjects:
            messagebox.showerror("خطا", "حداقل یک «موضوع» را از لیست انتخاب کنید.")
            switch_top("main")
            return

        if current_mode == "model":
            models = [m.strip() for m in models_text.get("1.0", "end").splitlines() if m.strip()]
            features = [f.strip() for f in feat1_text.get("1.0", "end").splitlines() if f.strip()]
            if not features:
                messagebox.showerror("خطا", "حداقل یک ویژگی وارد کنید (در تب «جستجو با مدل» اجباری است)."); return
            category = ""
            manual_features: List[str] = []
        else:
            models = []
            features = []
            category = category_var.get().strip()
            if manual_feat_var.get():
                manual_features = [f.strip() for f in feat2_text.get("1.0", "end").splitlines() if f.strip()]
                if not manual_features:
                    messagebox.showerror("خطا", "ویژگی‌های دستی را وارد کنید یا تیک «دستی وارد می‌کنم» را بردارید."); return
            else:
                manual_features = []

        save_config(config_path, {
            "gemini_api_keys": api_keys,
            "gemini_model": model_name,
            "last_inputs": inputs,
            "last_mode": current_mode,
            "last_models": models,
            "last_features": features,
            "last_category": category,
            "manual_features_enabled": manual_feat_var.get(),
            "last_manual_features": [f.strip() for f in feat2_text.get("1.0", "end").splitlines() if f.strip()],
            "last_subjects": subjects,
            "image_desc_enabled": desc_var.get(),
            "_subjects_cache": cfg.get("_subjects_cache", []),
        })

        runtime_opts["desc_enabled"] = bool(desc_var.get())
        cancel_event.clear()
        start_btn.config(state="disabled")
        cancel_btn.pack(side="left", padx=(10, 0))
        log_text.config(state="normal"); log_text.delete("1.0", "end"); log_text.config(state="disabled")
        log("شروع پردازش...")
        set_status("⏳ در حال پردازش...")

        if current_mode == "model":
            Thread(target=run_model_mode, args=(api_keys, model_name, inputs, models, features, subjects), daemon=True).start()
        else:
            Thread(target=run_auto_mode, args=(api_keys, model_name, inputs, category, manual_features, subjects), daemon=True).start()

    def on_cancel_extraction():
        cancel_event.set()
        set_status("⏳ در حال لغو...")

    def on_exit():
        root.destroy()

    start_btn = tk.Button(
        btn_frame, text="  شروع استخراج  ", command=on_start,
        bg=C_GREEN, fg=C_WHITE, font=("Arial", 11, "bold"),
        padx=20, pady=8, relief="flat", cursor="hand2",
        activebackground="#0daa6a", activeforeground=C_WHITE, bd=0,
    )
    start_btn.pack(side="left")

    cancel_btn = tk.Button(
        btn_frame, text="توقف", command=on_cancel_extraction,
        bg=C_RED, fg=C_WHITE, font=("Arial", 10),
        padx=12, pady=8, relief="flat", cursor="hand2",
        activebackground="#ee5555", bd=0,
    )

    tk.Button(
        btn_frame, text="خروج", command=on_exit,
        bg=C_GRAY, fg=C_WHITE, font=("Arial", 10),
        padx=12, pady=8, relief="flat", cursor="hand2",
        activebackground="#4a7a80", bd=0,
    ).pack(side="right")

    root.protocol("WM_DELETE_WINDOW", on_exit)
    root.mainloop()


# ======================== MAIN ========================

def main():
    app_dir = get_app_dir()

    if sys.stdout is not None:
        print("=" * 60)
        print("PDF Catalog Extractor")
        print("=" * 60)

    launch_ui(app_dir)


if __name__ == "__main__":
    try:
        main()
    except KeyboardInterrupt:
        if sys.stdout is not None:
            print("\n\n⚠ متوقف شد.")
    except Exception as e:
        if sys.stdout is not None:
            print(f"\n✗ خطای غیرمنتظره: {e}")
        logging.error(f"Fatal error: {e}", exc_info=True)
        raise
